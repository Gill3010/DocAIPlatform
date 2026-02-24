"""
PDF ↔ DOCX Conversion Converters
Uses pdf2docx for PDF->DOCX (preserves tables and images).
Docx->PDF: LibreOffice (host o Docker) primero, luego ReportLab con tablas.
"""
from pypdf import PdfReader

# Docker se importa solo cuando LibreOffice no está en el host
docker = None
from docx import Document
from docx.document import Document as _Document
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table as RlTable, TableStyle, PageBreak, Image as RlImage
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from typing import List, Iterator
import os
import io
import shutil
import tempfile
import subprocess
from pathlib import Path

from app.utils.base_converter import BaseConverter, ConversionError


def _pdf_to_docx_pdf2docx(input_path: str, output_path: str) -> bool:
    """Use pdf2docx library - preserves tables and images."""
    try:
        from pdf2docx import Converter
        cv = Converter(input_path)
        cv.convert(output_path, start=0, end=None)
        cv.close()
        return True
    except ImportError:
        return False
    except Exception:
        return False


def _pdf_to_docx_fallback(input_path: str, output_path: str) -> bool:
    """Fallback: PyMuPDF (mejor extracción) o pypdf + python-docx (text only, no tables/images)."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Prefer PyMuPDF for better layout-aware text extraction
    try:
        import fitz
        with fitz.open(input_path) as pdf_doc:
            if len(pdf_doc) == 0:
                raise ConversionError("PDF has no pages")
            for page_num in range(len(pdf_doc)):
                page = pdf_doc[page_num]
                blocks = page.get_text("blocks", sort=True)
                text_parts = []
                for b in blocks:
                    if len(b) >= 5 and b[4]:
                        text_parts.append(b[4].strip())
                text = '\n\n'.join(p for p in text_parts if p)
                if text:
                    if len(pdf_doc) > 1:
                        doc.add_heading(f'Página {page_num + 1}', level=2)
                    for para_text in text.split('\n\n'):
                        if para_text.strip():
                            doc.add_paragraph(para_text.strip())
                    if page_num < len(pdf_doc) - 1:
                        doc.add_page_break()
        doc.save(output_path)
        return True
    except ImportError:
        pass

    # Fallback to pypdf
    reader = PdfReader(input_path)
    if len(reader.pages) == 0:
        raise ConversionError("PDF has no pages")
    for page_num, page in enumerate(reader.pages, 1):
        text = page.extract_text()
        if text and text.strip():
            if len(reader.pages) > 1:
                doc.add_heading(f'Página {page_num}', level=2)
            for para_text in text.split('\n\n'):
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())
            if page_num < len(reader.pages):
                doc.add_page_break()
    doc.save(output_path)
    return True


class PDFToDocxConverter(BaseConverter):
    """PDF→DOCX via pdf2docx. Siempre local (prefers_local=True)."""
    
    @property
    def source_formats(self) -> List[str]:
        return ['pdf']
    
    @property
    def target_formats(self) -> List[str]:
        return ['docx']
    
    def convert(self, input_path: str, output_path: str) -> bool:
        self.ensure_directory(output_path)
        if _pdf_to_docx_pdf2docx(input_path, output_path):
            return True
        return _pdf_to_docx_fallback(input_path, output_path)


def _docx_to_pdf_libreoffice(input_path: str, output_path: str) -> bool:
    """Use LibreOffice headless - preserves tables and images. Returns True if successful."""
    libreoffice_cmd = shutil.which("libreoffice") or shutil.which("soffice")
    if not libreoffice_cmd:
        return False
    try:
        with tempfile.TemporaryDirectory(prefix="docx2pdf_") as tmpdir:
            result = subprocess.run(
                [libreoffice_cmd, "--headless", "--convert-to", "pdf", "--outdir", tmpdir, input_path],
                capture_output=True, text=True, timeout=120, env={**os.environ, "HOME": tmpdir},
            )
            if result.returncode != 0:
                return False
            out_pdf = Path(tmpdir) / f"{Path(input_path).stem}.pdf"
            if out_pdf.exists():
                shutil.copy2(out_pdf, output_path)
                return True
    except Exception:
        pass
    return False


def _iter_block_items(parent) -> Iterator:
    """Iterate paragraphs and tables in document order."""
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    else:
        return
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield DocxParagraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield DocxTable(child, parent)


def _docx_to_pdf_docker(input_path: str, output_path: str, base_name: str) -> bool:
    """Usa contenedor Docker document-converter para DOCX→PDF cuando LibreOffice no está en el host."""
    global docker
    if docker is None:
        try:
            import docker as docker_module
            docker = docker_module
        except ImportError:
            return False
    try:
        client = docker.from_env()
    except Exception:
        return False
    images = client.images.list(name="document-converter")
    if not images:
        images = client.images.list(name="766092484543.dkr.ecr.us-east-2.amazonaws.com/document-converter")
    if not images:
        return False
    image = images[0]
    input_dir = os.path.dirname(input_path)
    output_dir = os.path.dirname(output_path)
    input_file = os.path.basename(input_path)
    try:
        client.containers.run(
            image.id,
            command=[
                "libreoffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", "/tmp/output",
                f"/tmp/input/{input_file}",
            ],
            volumes={
                input_dir: {"bind": "/tmp/input", "mode": "ro"},
                output_dir: {"bind": "/tmp/output", "mode": "rw"},
            },
            remove=True,
        )
        expected = os.path.join(output_dir, f"{base_name}.pdf")
        if os.path.exists(expected):
            if expected != output_path:
                shutil.move(expected, output_path)
            return True
    except Exception:
        pass
    return False


class DocxToPDFConverter(BaseConverter):
    """Convert DOCX to PDF - LibreOffice (host o Docker), luego ReportLab con tablas."""

    @property
    def prefers_local(self) -> bool:
        """Permite ECS cuando USE_ECS_CONVERTER=true (mejor calidad tablas/imágenes)."""
        return False

    @property
    def source_formats(self) -> List[str]:
        return ['docx']

    @property
    def target_formats(self) -> List[str]:
        return ['pdf']

    def convert(self, input_path: str, output_path: str) -> bool:
        self.ensure_directory(output_path)
        input_path = os.path.abspath(input_path)
        output_path = os.path.abspath(output_path)
        base_name = Path(input_path).stem
        if _docx_to_pdf_libreoffice(input_path, output_path):
            return True
        if _docx_to_pdf_docker(input_path, output_path, base_name):
            return True
        return self._convert_with_reportlab(input_path, output_path)
    
    def _convert_with_reportlab(self, input_path: str, output_path: str) -> bool:
        """ReportLab conversion with paragraphs AND tables in order."""
        doc = Document(input_path)
        pdf = SimpleDocTemplate(output_path, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
        styles = getSampleStyleSheet()
        normal_style, heading_style = styles['Normal'], styles['Heading1']
        elements = []
        for block in _iter_block_items(doc):
            if isinstance(block, DocxParagraph):
                # Add inline images first (in document order)
                for run in block.runs:
                    try:
                        for child in run._element.iterdescendants():
                            if "blip" in str(child.tag).lower():
                                rId = child.get(qn("r:embed")) or child.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                                if rId and hasattr(doc.part, "related_parts") and rId in doc.part.related_parts:
                                    img_part = doc.part.related_parts[rId]
                                    ext = ".png" if "png" in str(img_part.content_type) else ".jpg"
                                    with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
                                        tmp.write(img_part.blob)
                                        tmp.flush()
                                        try:
                                            img = RlImage(tmp.name, width=300, height=200)
                                            elements.append(Spacer(1, 12))
                                            elements.append(img)
                                            elements.append(Spacer(1, 12))
                                        finally:
                                            os.unlink(tmp.name)
                                break
                    except Exception:
                        pass
                text = block.text.strip()
                if text:
                    style = heading_style if block.style.name.startswith('Heading') else normal_style
                    elements.append(Paragraph(text, style))
                    elements.append(Spacer(1, 12))
            elif isinstance(block, DocxTable):
                rows = []
                max_cols = max(len(row.cells) for row in block.rows) if block.rows else 0
                for row in block.rows:
                    cells = [str(cell.text or "").strip()[:80] for cell in row.cells]
                    cells.extend([""] * (max_cols - len(cells)))
                    rows.append(cells)
                if rows:
                    t = RlTable(rows)
                    t.setStyle(TableStyle([
                        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 0), (-1, -1), 8),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                    ]))
                    elements.append(Spacer(1, 12))
                    elements.append(t)
                    elements.append(Spacer(1, 12))
        if not elements:
            raise ConversionError("El documento DOCX está vacío")
        pdf.build(elements)
        return True
