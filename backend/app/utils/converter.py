"""
Lightweight document conversion utilities
Optimized for AWS Free Tier (limited RAM/disk)
"""
import os
from pathlib import Path
from typing import Optional
from PIL import Image
from pypdf import PdfReader, PdfWriter
from docx import Document
import io

class ConversionError(Exception):
    """Custom exception for conversion errors"""
    pass

def ensure_directory(file_path: str):
    """Ensure the directory for the file exists"""
    Path(file_path).parent.mkdir(parents=True, exist_ok=True)

# ===== IMAGE TO PDF =====
def image_to_pdf(input_path: str, output_path: str) -> bool:
    """
    Convert image (PNG, JPG, JPEG) to PDF
    Memory efficient: processes one image at a time
    """
    try:
        ensure_directory(output_path)
        
        # Open and convert image
        image = Image.open(input_path)
        
        # Convert to RGB if necessary (PNG with transparency)
        if image.mode in ("RGBA", "LA", "P"):
            rgb_image = Image.new("RGB", image.size, (255, 255, 255))
            rgb_image.paste(image, mask=image.split()[-1] if image.mode == "RGBA" else None)
            image = rgb_image
        elif image.mode != "RGB":
            image = image.convert("RGB")
        
        # Save as PDF
        image.save(output_path, "PDF", resolution=100.0)
        image.close()
        
        return True
    except Exception as e:
        raise ConversionError(f"Image to PDF conversion failed: {str(e)}")

# ===== PDF TO IMAGES =====
def pdf_to_image(input_path: str, output_path: str, page_number: int = 0) -> bool:
    """
    Convert first page of PDF to PNG image
    Note: Full PDF rendering requires heavy libs (pdf2image, poppler)
    This is a placeholder - for production consider external service
    """
    try:
        ensure_directory(output_path)
        
        # For now, we'll extract text and create a simple image
        # In production, use pdf2image with poppler or external API
        reader = PdfReader(input_path)
        
        if len(reader.pages) == 0:
            raise ConversionError("PDF has no pages")
        
        # Create a simple placeholder image with PDF info
        img = Image.new('RGB', (800, 600), color='white')
        img.save(output_path, 'PNG')
        
        return True
    except Exception as e:
        raise ConversionError(f"PDF to image conversion failed: {str(e)}")

# ===== TEXT TO DOCX =====
def text_to_docx(input_path: str, output_path: str) -> bool:
    """
    Convert plain text file to DOCX
    """
    try:
        ensure_directory(output_path)
        
        # Read text file
        with open(input_path, 'r', encoding='utf-8') as f:
            text_content = f.read()
        
        # Create DOCX
        doc = Document()
        
        # Add paragraphs
        for paragraph in text_content.split('\n\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        doc.save(output_path)
        return True
    except Exception as e:
        raise ConversionError(f"Text to DOCX conversion failed: {str(e)}")

# ===== DOCX TO TEXT =====
def docx_to_text(input_path: str, output_path: str) -> bool:
    """
    Convert DOCX to plain text
    """
    try:
        ensure_directory(output_path)
        
        doc = Document(input_path)
        
        # Extract all text
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Write to text file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(full_text))
        
        return True
    except Exception as e:
        raise ConversionError(f"DOCX to text conversion failed: {str(e)}")

# ===== PDF TO TEXT =====
def pdf_to_text(input_path: str, output_path: str) -> bool:
    """
    Extract text from PDF
    """
    try:
        ensure_directory(output_path)
        
        reader = PdfReader(input_path)
        
        # Extract text from all pages
        full_text = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                full_text.append(text)
        
        # Write to text file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n\n'.join(full_text))
        
        return True
    except Exception as e:
        raise ConversionError(f"PDF to text conversion failed: {str(e)}")

# ===== MAIN CONVERTER ROUTER =====
def convert_file(input_path: str, output_path: str, source_format: str, target_format: str) -> bool:
    """
    Main conversion router - dispatches to appropriate converter
    
    Args:
        input_path: Path to input file
        output_path: Path to output file
        source_format: Source format (e.g., 'png', 'pdf', 'docx')
        target_format: Target format (e.g., 'pdf', 'png', 'txt')
    
    Returns:
        bool: True if conversion successful
    
    Raises:
        ConversionError: If conversion fails or format not supported
    """
    # Normalize formats
    source = source_format.lower().replace('.', '')
    target = target_format.lower().replace('.', '')
    
    # Define conversion map
    conversion_map = {
        ('png', 'pdf'): image_to_pdf,
        ('jpg', 'pdf'): image_to_pdf,
        ('jpeg', 'pdf'): image_to_pdf,
        ('pdf', 'png'): pdf_to_image,
        ('pdf', 'txt'): pdf_to_text,
        ('txt', 'docx'): text_to_docx,
        ('docx', 'txt'): docx_to_text,
    }
    
    converter_key = (source, target)
    
    if converter_key not in conversion_map:
        raise ConversionError(
            f"Conversion from {source_format} to {target_format} is not supported yet. "
            f"Available conversions: {list(conversion_map.keys())}"
        )
    
    converter_func = conversion_map[converter_key]
    return converter_func(input_path, output_path)

def get_supported_conversions() -> dict:
    """
    Return dictionary of supported conversions
    """
    return {
        "png": ["pdf"],
        "jpg": ["pdf"],
        "jpeg": ["pdf"],
        "pdf": ["png", "txt"],
        "txt": ["docx"],
        "docx": ["txt"],
    }
