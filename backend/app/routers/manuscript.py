import re
import io
import unicodedata
from datetime import datetime
from pathlib import Path
from typing import Optional, Union, BinaryIO, List

import urllib.parse
from fastapi import APIRouter, Depends, UploadFile, File, Form, HTTPException, status
from fastapi.responses import StreamingResponse, FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.core.database import get_db
from app.core.security import get_current_user
from app.models.user import User
from app.models.manuscript_format import ManuscriptFormat
from app.schemas.manuscript_format import ManuscriptFormatResponse

from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Cm, Pt

router = APIRouter()

# Storage directory for formatted manuscripts
_BASE_DIR = Path(__file__).resolve().parent.parent.parent  # backend/
FORMATTED_DIR = _BASE_DIR / "storage" / "formatted"
FORMATTED_DIR.mkdir(parents=True, exist_ok=True)


def _safe_filename(name: str) -> str:
    """Normalize filename to be safe for filesystem and HTTP headers."""
    n = unicodedata.normalize("NFC", name)
    try:
        n.encode("latin-1")
        return n
    except UnicodeEncodeError:
        return n.encode("ascii", "replace").decode("ascii").replace("?", "_")


class ManuscriptFormatter:
    """
    Format unstructured academic manuscripts (.docx) into a standardized
    structure compatible with the JATS XML converter.

    Normalization rules applied:
    - Front matter: Title (centered+bold), Trans-title (centered+italic),
      Author block (right-aligned), Dates/URL/DOI (left-aligned).
    - Filler label paragraphs (Autora, Institución, Correspondencia, etc.) removed.
    - Body section headings normalized to standard labels + Heading 1 style.
    - References section normalized to 'Referencias' + Heading 1 style.
    - Front matter keywords/abstract labels normalized + Heading 2 style.
    - Figure/Table captions standardized (Gráfico→Figura, Cuadro→Tabla).
    - Double spaces collapsed.
    - Excessive blank paragraphs removed.
    """

    BODY_SECTION_MAPPINGS = {
        r'(?i)^\s*(?:\d+(\.\d+)*|[ivx]+|uno|dos|tres|cuatro|cinco|seis)\.?\s*(introducci[oó]n|introduction|intro)\s*$': 'Introducción',
        r'(?i)^\s*(?:\d+(\.\d+)*|[ivx]+|uno|dos|tres|cuatro|cinco|seis)\.?\s*(metodolog[ií]a|methodology|m[eé]todos?|methods)\s*$': 'Metodología',
        r'(?i)^\s*(?:\d+(\.\d+)*|[ivx]+|uno|dos|tres|cuatro|cinco|seis)\.?\s*(resultados|results)\s*$': 'Resultados',
        r'(?i)^\s*(?:\d+(\.\d+)*|[ivx]+|uno|dos|tres|cuatro|cinco|seis)\.?\s*(discusi[oó]n|discussion)\s*$': 'Discusión',
        r'(?i)^\s*(?:\d+(\.\d+)*|[ivx]+|uno|dos|tres|cuatro|cinco|seis)\.?\s*(conclusiones|conclusion|conclusi[oó]n|conclusions)\s*$': 'Conclusiones',
        r'(?i)^\s*(?:\d+(\.\d+)*|[ivx]+|uno|dos|tres|cuatro|cinco|seis)\.?\s*(marco te[oó]rico|antecedentes|revisi[oó]n bibliogr[aá]fica)\s*$': 'Marco Teórico',
    }

    REFS_MAPPINGS = {
        r'(?i)^\s*(?:\d+(\.\d+)*|[ivx]+|uno|dos|tres|cuatro|cinco|seis)\.?\s*(referencias(?: bibliogr[aá]ficas)?|bibliograf[ií]a|references|bibliography)\s*$': 'Referencias'
    }

    FRONT_MATTER_MAPPINGS = {
        r'(?i)^\s*(resumen)\s*$': 'Resumen',
        r'(?i)^\s*(abstract)\s*$': 'Abstract',
        r'(?i)^\s*(palabras\s*clave[s]?|palabras-clave)\s*$': 'Palabras clave',
        r'(?i)^\s*(keywords|key\s*words)\s*$': 'Keywords'
    }

    # Labels that are pure decorative headers in unstructured docs — remove them
    FILLER_LABELS = {
        "autora", "autor", "autores",
        "institución", "instituciones", "institucion", "instituciones",
        "filiación", "filiacion", "afiliación", "afiliacion",
        "correspondencia:", "correspondencia",
        "contacto:", "contacto",
        "orcid:", "orcid",
    }

    def format_manuscript(self, input_file: Union[str, BinaryIO, bytes]) -> io.BytesIO:
        """
        Receives a docx file path, bytes, or file-like object and returns a formatted BytesIO.
        """
        if isinstance(input_file, bytes):
            doc = Document(io.BytesIO(input_file))
        else:
            doc = Document(input_file)

        self._normalize_document(doc)

        output_stream = io.BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)
        return output_stream

    def _normalize_document(self, doc: Document):
        """Applies normalization rules to the document."""
        self._reorder_front_matter(doc)
        self._apply_global_apa_format(doc)
        self._normalize_paragraphs(doc)
        self._format_references(doc)
        self._clean_empty_paragraphs(doc)

    def _apply_global_apa_format(self, doc: Document):
        """Sets standard APA document-level margins."""
        for section in doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

    def _reorder_front_matter(self, doc: Document):
        """Finds the translated title (if misplaced before Abstract) and moves it below main title."""
        paras = [p for p in doc.paragraphs if p.text.strip() and not p.text.strip().lower() in self.FILLER_LABELS]
        if not paras: return

        main_title = paras[0]

        for i, p in enumerate(paras):
            text = p.text.strip().lower()
            if text in ('abstract', 'abstract:'):
                if i > 1:
                    candidate = paras[i - 1]
                    ctext = candidate.text.strip().lower()
                    # Ensure it's not keywords, nor author blocks
                    if not any(k in ctext for k in ('palabras clave', 'keywords', 'resumen', '@', 'orcid', 'universidad', 'university', 'departamento')):
                        # It is the English Title. Move it.
                        elem = candidate._element
                        elem.getparent().remove(elem)
                        main_title._element.addnext(elem)
                break

    def _safe_set_style(self, para: Paragraph, style_name: str):
        """Safely set a paragraph style, falling back if not found."""
        try:
            para.style = style_name
        except KeyError:
            try:
                if style_name == 'Heading 1':
                    para.style = 'Título 1'
                elif style_name == 'Heading 2':
                    para.style = 'Título 2'
            except KeyError:
                para.style = 'Normal'

    def _replace_paragraph_text(self, paragraph: Paragraph, new_text: str):
        """Replaces text in a paragraph keeping it as simple text."""
        paragraph.text = new_text

    def _normalize_captions(self, text: str) -> str:
        """Standardize caption labels (e.g., Gráfico -> Figura, Cuadro -> Tabla)."""
        text = re.sub(r'(?i)^(gr[aá]fico|ilustraci[oó]n|imagen)\s+(\d+)', r'Figura \2', text)
        text = re.sub(r'(?i)^(cuadro)\s+(\d+)', r'Tabla \2', text)
        return text

    def _detect_section(self, text: str, mappings: dict) -> str:
        """Helper to match paragraph text against a set of regex patterns."""
        for pattern, standard_name in mappings.items():
            if re.match(pattern, text):
                return standard_name
        return ""

    def _normalize_paragraphs(self, doc: Document):
        """Iterates through paragraphs and applies structural rules."""
        is_body_started = False
        is_references = False
        has_abstract = False
        non_empty_count = 0

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            non_empty_count += 1
            lower_text = text.lower()
            style_name = getattr(para.style, 'name', '')

            # Global paragraph baseline formatting (APA)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.space_after = Cm(0)
            para.paragraph_format.space_before = Cm(0)
            
            is_header_style = style_name.startswith(('Heading', 'Título', 'Title', 'Subtitle'))
            for run in para.runs:
                run.font.name = 'Times New Roman'
                # Enforce body font size if it has explicit overrides or is normal text
                if not is_header_style:
                    run.font.size = Pt(12)

            # 1. Check for References start
            matched_refs = self._detect_section(text, self.REFS_MAPPINGS)
            if matched_refs:
                self._replace_paragraph_text(para, matched_refs)
                self._safe_set_style(para, 'Heading 1')
                is_references = True
                continue

            # 2. Check for Body Sections
            if not is_references:
                matched_body = self._detect_section(text, self.BODY_SECTION_MAPPINGS)
                if matched_body:
                    self._replace_paragraph_text(para, matched_body)
                    self._safe_set_style(para, 'Heading 1')
                    para.paragraph_format.first_line_indent = Cm(0)
                    para.alignment = 0 # Left
                    is_body_started = True
                    continue

            # 3. Check for Front Matter Headers (Resumen, Abstract, Keywords)
            if not is_body_started and not is_references:
                matched_front = self._detect_section(text, self.FRONT_MATTER_MAPPINGS)
                if matched_front:
                    self._replace_paragraph_text(para, matched_front)
                    self._safe_set_style(para, 'Heading 2')
                    para.alignment = 0  # Left
                    para.paragraph_format.first_line_indent = Cm(0)
                    has_abstract = True
                    continue

            # 4. Front Matter block (before body and before abstract headers)
            if not is_body_started and not is_references and not has_abstract:
                # Dates, URL, DOI → Left align, keep as-is
                if lower_text.startswith(("recibido", "aceptado", "url:", "doi:", "url ", "doi ")):
                    para.alignment = 0  # Left
                    para.paragraph_format.first_line_indent = Cm(0)
                    continue

                # Remove filler/decorator labels
                clean_lower = lower_text.strip().rstrip(":")
                if clean_lower in self.FILLER_LABELS or (len(text) < 35 and clean_lower in self.FILLER_LABELS):
                    para.clear()
                    non_empty_count -= 1
                    continue

                # Title (first substantive paragraph) → centered + bold
                if non_empty_count == 1:
                    self._safe_set_style(para, 'Title')
                    para.alignment = 1  # Center
                    para.paragraph_format.first_line_indent = Cm(0)
                    for run in para.runs:
                        run.bold = True
                    continue

                # Trans-title (second substantive paragraph, long enough, no email/orcid)
                if non_empty_count == 2 and len(text) > 30 and "@" not in text and "orcid" not in lower_text:
                    para.alignment = 1  # Center
                    para.paragraph_format.first_line_indent = Cm(0)
                    for run in para.runs:
                        run.italic = True
                    continue

                # Author block → Right aligned
                para.alignment = 2  # Right
                para.paragraph_format.first_line_indent = Cm(0)
                continue

            # 5. Clean up captions (Figures, Tables)
            is_caption = False
            if re.match(r'(?i)^(gr[aá]fico|ilustraci[oó]n|imagen|cuadro|figura|tabla)\s+\d+', text):
                new_text = self._normalize_captions(text)
                self._replace_paragraph_text(para, new_text)
                text = new_text
                is_caption = True
                para.paragraph_format.first_line_indent = Cm(0)
                para.alignment = 1 # Center

            # Body text indentation
            if is_body_started and not is_references and not is_caption:
                # Regular paragraph in body
                para.alignment = 3 # Justified
                para.paragraph_format.first_line_indent = Cm(1.27)
            elif has_abstract and not is_body_started and not is_references and not is_caption:
                # Abstract text
                para.alignment = 3 # Justified
                para.paragraph_format.first_line_indent = Cm(0)
                if text.lower().startswith(("palabras", "keywords")):
                    para.paragraph_format.first_line_indent = Cm(1.27)

            # 6. Clean up multiple spaces
            clean_text = re.sub(r' {2,}', ' ', para.text)
            if clean_text != para.text:
                self._replace_paragraph_text(para, clean_text)

    def _format_references(self, doc: Document):
        """Consolidates reference fragments and applies APA hanging indent."""
        in_refs = False
        ref_paras = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not in_refs:
                if text.lower() == 'referencias':
                    in_refs = True
                continue
                
            if not text:
                continue
            ref_paras.append(para)

        if not ref_paras:
            return

        current_para = None

        for para in ref_paras:
            text = para.text.strip()
            
            # Simple heuristic: it's a continuation if it starts with lower case or 'http', or if previous didn't end with a period.
            is_continuation = text.lower().startswith('http') or text[0].islower()
            
            if current_para is not None and is_continuation:
                if current_para.text and not current_para.text.endswith(' ') and text and not text.startswith(' '):
                    current_para.add_run(" ")
                
                for run in para.runs:
                    new_run = current_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                
                p = para._element
                p.getparent().remove(p)
            else:
                if current_para is not None:
                    self._apply_apa_format(current_para)
                current_para = para

        if current_para is not None:
            self._apply_apa_format(current_para)

    def _apply_apa_format(self, para: Paragraph):
        """Applies APA 7th edition formatting to a reference paragraph."""
        para.paragraph_format.left_indent = Cm(1.27)
        para.paragraph_format.first_line_indent = Cm(-1.27)
        para.alignment = 3  # WD_ALIGN_PARAGRAPH.JUSTIFY

    def _clean_empty_paragraphs(self, doc: Document):
        """Removes all empty paragraphs by completely deleting their XML element."""
        for para in doc.paragraphs:
            if not para.text.strip():
                p = para._element
                p.getparent().remove(p)


# ─────────────────────────────────────────────
# Helper: access check
# ─────────────────────────────────────────────

def _check_pro_access(user: User):
    is_pro = (
        user.is_superuser
        or getattr(user, "can_access_admin_panel", False)
        or (user.is_premium and user.premium_plan_id in ("Pro", "Empresa"))
    )
    if not is_pro:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="pro_plan_required")


# ─────────────────────────────────────────────
# Endpoints
# ─────────────────────────────────────────────

@router.post(
    "/format",
    summary="Formatear manuscrito (Solo Pro/Empresa)",
    description=(
        "Aplica formato profesional a un documento. "
        "El archivo formateado se guarda y queda disponible en el historial del usuario. "
        "Requiere plan Pro o Empresa."
    ),
    responses={
        200: {"description": "Documento formateado y descargado"},
        400: {"description": "Extensión de archivo no soportada"},
        403: {"description": "Acceso restringido (se requiere plan Pro o superior)"},
        500: {"description": "Error al procesar el archivo"},
    },
)
async def format_manuscript(
    file: UploadFile = File(...),
    style: str = Form("standard"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    _check_pro_access(current_user)

    if not (file.filename or "").endswith(".docx"):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Solamente se soportan documentos en formato .docx",
        )

    content = await file.read()
    file_size_mb = round(len(content) / (1024 * 1024), 4)
    original_filename = file.filename or "documento.docx"

    # Build output filename and path
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = re.sub(r'[^\w.\-]', '_', Path(original_filename).stem)
    output_filename = f"user_{current_user.id}_{timestamp}_{safe_name}_formatted.docx"
    output_path = FORMATTED_DIR / output_filename

    # Create DB record (before processing, to have an id)
    record = ManuscriptFormat(
        user_id=current_user.id,
        original_filename=original_filename,
        file_size=file_size_mb,
        status="processing",
    )
    db.add(record)
    await db.commit()
    await db.refresh(record)

    try:
        formatter = ManuscriptFormatter()
        output_stream = formatter.format_manuscript(content)

        # Save to disk
        with open(output_path, "wb") as f:
            f.write(output_stream.read())

        # Update record
        record.output_file_path = str(output_path)
        record.status = "completed"
        db.add(record)
        await db.commit()

        output_stream.seek(0)

    except Exception as e:
        record.status = "failed"
        record.error_message = str(e)
        db.add(record)
        await db.commit()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error durante el formateo del manuscrito: {str(e)}",
        )

    safe_filename = urllib.parse.quote(original_filename)
    headers = {
        "Content-Disposition": f"attachment; filename*=utf-8''formateado_{safe_filename}"
    }
    output_stream = io.BytesIO(open(output_path, "rb").read())
    return StreamingResponse(
        output_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


@router.get(
    "/history",
    response_model=List[ManuscriptFormatResponse],
    summary="Historial de manuscritos formateados",
    description="Lista los manuscritos formateados del usuario autenticado, ordenados por fecha descendente.",
)
async def get_manuscript_history(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
    limit: int = 20,
):
    _check_pro_access(current_user)
    result = await db.execute(
        select(ManuscriptFormat)
        .where(ManuscriptFormat.user_id == current_user.id)
        .order_by(ManuscriptFormat.created_at.desc())
        .limit(limit)
    )
    return result.scalars().all()


@router.get(
    "/download/{format_id}",
    summary="Descargar manuscrito formateado",
    description="Descarga el archivo formateado previamente. Solo disponible para el usuario propietario.",
    responses={
        200: {"description": "Archivo .docx formateado"},
        403: {"description": "Acceso restringido"},
        404: {"description": "Registro no encontrado o archivo eliminado"},
    },
)
async def download_formatted_manuscript(
    format_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    _check_pro_access(current_user)

    result = await db.execute(
        select(ManuscriptFormat).where(
            ManuscriptFormat.id == format_id,
            ManuscriptFormat.user_id == current_user.id,
        )
    )
    record = result.scalar_one_or_none()

    if not record:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Registro no encontrado")

    if record.status != "completed" or not record.output_file_path:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"El formateo no está completado. Estado: {record.status}",
        )

    output_path = Path(record.output_file_path)
    if not output_path.exists():
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="El archivo ya no existe en el servidor",
        )

    download_name = _safe_filename(f"formateado_{record.original_filename}")
    return FileResponse(
        path=str(output_path),
        filename=download_name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
