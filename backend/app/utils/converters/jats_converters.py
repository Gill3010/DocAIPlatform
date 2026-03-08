"""
DOCX ↔ XML JATS Conversion Converters
Handles: Microsoft Word to JATS (Journal Article Tag Suite) XML format
Used for academic/scientific article publishing.
Generates external media references for OJS dependent files.
"""
from docx import Document
from docx.table import Table
from lxml import etree
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Union
import re
import unicodedata
from datetime import datetime

from app.core.config import settings
from app.utils.base_converter import BaseConverter, ConversionError


class DocxToJATSConverter(BaseConverter):
    """Convert Microsoft Word document to JATS XML format."""

    XLINK_NS = 'http://www.w3.org/1999/xlink'
    XML_NS = 'http://www.w3.org/XML/1998/namespace'

    # Compiled patterns
    _ORCID_RE = re.compile(
        r'(?:https?://orcid\.org/)?(\d{4}-\d{4}-\d{4}-\d{3}[\dX])', re.IGNORECASE
    )
    _EMAIL_RE = re.compile(r'\b[\w.+\-]+@[\w.\-]+\.[a-zA-Z]{2,}\b')
    _DOI_RE = re.compile(
        # Handles: "DOI: 10.xxx", "DOI: https://doi.org/10.xxx", "https://doi.org/10.xxx"
        r'(?:DOI\s*[:\s]+\s*(?:https?://(?:dx\.)?doi\.org/)?|https?://(?:dx\.)?doi\.org/)([\w./\-]+)',
        re.IGNORECASE,
    )
    _URL_RE = re.compile(r'https?://[^\s<>"\']+')
    _YEAR_RE = re.compile(r'\b(19|20)\d{2}\b')
    _DATE_ES_RE = re.compile(r'(\d{1,2})\s+de\s+(\w+)\s+de\s+(\d{4})', re.IGNORECASE)
    _DATE_SLASH_RE = re.compile(r'(\d{1,2})[/.\-](\d{1,2})[/.\-](\d{4})')
    _MONTHS_ES: Dict[str, int] = {
        'enero': 1, 'febrero': 2, 'marzo': 3, 'abril': 4, 'mayo': 5, 'junio': 6,
        'julio': 7, 'agosto': 8, 'septiembre': 9, 'octubre': 10, 'noviembre': 11,
        'diciembre': 12, 'january': 1, 'february': 2, 'march': 3, 'april': 4,
        'may': 5, 'june': 6, 'july': 7, 'august': 8, 'september': 9,
        'october': 10, 'november': 11, 'december': 12,
    }

    # Section headings that mark the start of the article body
    _BODY_SECTION_LABELS = {
        'introducción', 'introduction', 'introduccion',
        'metodología', 'methodology', 'metodos', 'methods', 'método',
        'resultados', 'results',
        'conclusión', 'conclusion', 'conclusiones', 'conclusions',
        'discusión', 'discussion', 'discusion',
        'antecedentes', 'marco teórico', 'marco teorico',
        'revisión de literatura', 'revisión bibliográfica',
        'desarrollo', 'hallazgos', 'findings',
    }

    # Headings that mark the references section (stop body extraction here)
    _REFS_LABELS = {
        'referencias', 'referencias bibliográficas', 'referencias bibliograficas',
        'bibliografía', 'bibliografia', 'references', 'bibliography',
        'fuentes bibliográficas', 'fuentes bibliograficas',
    }

    # Country names used to detect country lines
    _COUNTRY_NAMES = {
        'perú', 'peru', 'colombia', 'argentina', 'mexico', 'méxico', 'chile',
        'venezuela', 'ecuador', 'bolivia', 'cuba', 'costa rica', 'panamá', 'panama',
        'honduras', 'nicaragua', 'el salvador', 'guatemala', 'dominicana',
        'república dominicana', 'españa', 'spain', 'brasil', 'brazil',
    }

    # Institution keywords
    _INST_KEYWORDS = (
        'universidad', 'universit', 'institute', 'college', 'escuela',
        'faculty', 'departament', 'depart', 'institución', 'centro de',
        'instituto', 'politécnico', 'tecnológico',
    )

    # Article-type labels that can appear before the actual title
    _ARTICLE_TYPES = {
        'artículo original', 'articulo original', 'original article',
        'artículo de revisión', 'articulo de revision', 'review article',
        'artículo de investigación', 'articulo de investigacion', 'research article',
        'artículo de reflexión', 'articulo de reflexion', 'reflection article',
        'reporte de caso', 'case report',
        'carta al editor', 'letter to the editor',
        'editorial', 'nota técnica', 'nota tecnica', 'short communication',
    }

    # In-text citation patterns
    # Pattern A: Author (YYYY) or Author y Author (YYYY)  → xref wraps only the year
    # Pattern B: (Author..., YYYY) or (Author... YYYY)    → xref wraps full inner content
    # Pattern A – Author (YYYY): handles "Apellido (año)", "Apellido, (año)",
    #   "A y B (año)", "A, B, & C, (año)", "A et al. (año)", etc.
    # Pattern B – (Author … YYYY): parenthetical citation, e.g. (García, 2020)
    _CITATION_RE = re.compile(
        r'([A-ZÁÉÍÓÚÑ][a-záéíóúñ]+'
        r'(?:,?\s+et\s+al\.?)?'
        r'(?:(?:,\s+(?:&\s+)?|\s+(?:y|and|&)\s+)[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+'
        r'(?:,?\s+et\s+al\.?)?)*)'
        r'[,\s]*\((\d{4}[a-z]?)(?:[^)]{0,30})?\)'
        r'|'
        r'\(([A-ZÁÉÍÓÚÑ][^)]*?)\s+(\d{4}[a-z]?)(?:[^)]{0,20})?\)'
    )

    # Figure/table caption and in-text reference patterns
    _FIG_CAPTION_RE = re.compile(r'^(Figura|Fig\.?)\s+(\d+[a-z]?)\s*[.\s:]*(.*)?$', re.IGNORECASE)
    _TABLE_CAPTION_RE = re.compile(r'^(Tabla|Tab\.?)\s+(\d+[a-z]?)\s*[.\s:]*(.*)?$', re.IGNORECASE)
    # Single reference (used for xref replacement in paragraph text)
    _FIGREF_RE = re.compile(
        r'\b(Figura|Fig\.?|Tabla|Tab\.?)\s*(?:N(?:[°º]|o\.?)\s*)?(\d+[a-z]?)\b',
        re.IGNORECASE,
    )
    # Wider matcher (used to collect all referenced media ids, including plurals/lists/ranges)
    _MEDIA_REF_BLOCK_RE = re.compile(
        r'\b(Figuras?|Figs?\.?|Tablas?|Tabs?\.?)\s*'
        r'(?:N(?:[°º]|o\.?)\s*)?'
        r'((?:\d+[a-z]?)(?:\s*(?:,|;|y|e|and|&|-|–)\s*\d+[a-z]?){0,30})',
        re.IGNORECASE,
    )

    @property
    def source_formats(self) -> List[str]:
        return ['docx']

    @property
    def target_formats(self) -> List[str]:
        return ['xml']

    # ------------------------------------------------------------------ #
    # Main entry point
    # ------------------------------------------------------------------ #

    def convert(self, input_path: str, output_path: str, journal_meta: Optional[Dict] = None) -> bool:
        """
        Convert DOCX to JATS XML.
        - Detects in-text references (Figura/Tabla N).
        - Creates external media references with forced names:
          fig-N.jpg and table-N.jpg.
        - Does not extract or package image files.
        - journal_meta: opcional dict con journal_id, journal_title, issn, publisher_name (multi-revista).
        """
        try:
            self.ensure_directory(output_path)

            doc = Document(input_path)

            metadata = self._extract_metadata(doc, input_path)
            body_sections = self._extract_body(doc)
            self._inject_external_media_references(doc, body_sections)
            references = self._extract_references(doc)

            jats_xml = self._build_jats_xml(metadata, body_sections, references, journal_meta=journal_meta)

            tree = etree.ElementTree(jats_xml)
            tree.write(
                output_path,
                pretty_print=True,
                xml_declaration=True,
                encoding='utf-8',
                doctype=(
                    '<!DOCTYPE article PUBLIC "-//NLM//DTD JATS (Z39.96) Journal Publishing DTD'
                    ' v1.1 20151215//EN"'
                    ' "https://jats.nlm.nih.gov/publishing/1.1/JATS-journalpublishing1.dtd">'
                ),
            )

            return True

        except Exception as e:
            raise ConversionError(f"Conversión DOCX a JATS XML falló: {str(e)}")

    # ------------------------------------------------------------------ #
    # Helpers
    # ------------------------------------------------------------------ #

    def _parse_date(self, text: str) -> Optional[Dict[str, str]]:
        """Return {'day', 'month', 'year'} parsed from text, or None."""
        m = self._DATE_ES_RE.search(text)
        if m:
            month_num = self._MONTHS_ES.get(m.group(2).lower())
            if month_num:
                return {
                    'day': m.group(1).zfill(2),
                    'month': str(month_num).zfill(2),
                    'year': m.group(3),
                }
        m2 = self._DATE_SLASH_RE.search(text)
        if m2:
            return {
                'day': m2.group(1).zfill(2),
                'month': m2.group(2).zfill(2),
                'year': m2.group(3),
            }
        return None

    def _parse_author_name(self, name_raw: str) -> Tuple[str, str]:
        """Split raw name string into (surname, given_names)."""
        name = self._ORCID_RE.sub('', name_raw)
        name = self._EMAIL_RE.sub('', name)
        name = re.sub(r'\s*[\*¹²³⁴⁵⁶⁷⁸⁹0-9]+\s*', ' ', name)
        name = re.sub(r'[¹²³⁴⁵⁶⁷⁸⁹]', '', name)
        parts = [p for p in name.split() if p and not re.match(r'^[\*.,]+$', p)]
        if not parts:
            return 'Autor', ''
        if len(parts) >= 4:
            return ' '.join(parts[-2:]), ' '.join(parts[:-2])
        if len(parts) == 3:
            return ' '.join(parts[-2:]), parts[0]
        if len(parts) == 2:
            return parts[-1], parts[0]
        return parts[0], ''

    def _deduplicate_text(self, text: str) -> str:
        """Remove exact internal duplication (common in some Word paragraph styles).

        Handles cases where text = A + A regardless of whether len(text) is
        perfectly even, by searching for the first 30 chars of the text as a
        repeated substring and verifying the two halves are identical.
        """
        n = len(text)
        if n < 20:
            return text
        # Try anchor-based search: take a prefix and look for it repeating
        prefix_len = min(30, n // 3)
        prefix = text[:prefix_len]
        pos = text.find(prefix, n // 3)
        if pos > 0 and text[pos:] == text[:pos]:
            return text[:pos]
        # Fallback: brute-force near the midpoint (handles tiny off-by-one cases)
        for split in range(n // 2 - 3, n // 2 + 4):
            if 0 < split < n and text[:split] == text[split:]:
                return text[:split]
        return text

    def _consolidate_references(self, parts: List[str]) -> List[str]:
        """Join multi-paragraph reference fragments into single entries."""
        if not parts:
            return []
        # A new reference starts with: optional number + Capital letter + comma/capital word
        new_ref_re = re.compile(
            r'^(?:\d+[\.\)]\s+)?[A-ZÁÉÍÓÚÑ][a-záéíóúñA-Z\-]*(?:,|\s+[A-ZÁÉÍÓÚÑ])'
        )
        consolidated: List[str] = []
        current: Optional[str] = None
        for part in parts:
            if new_ref_re.match(part):
                if current is not None:
                    consolidated.append(current)
                current = part
            else:
                current = (current + ' ' + part) if current is not None else part
        if current is not None:
            consolidated.append(current)
        return consolidated

    # ------------------------------------------------------------------ #
    # Front-matter extraction
    # ------------------------------------------------------------------ #

    def _extract_metadata(self, doc: Document, input_path: Optional[str] = None) -> Dict:
        """
        Comprehensive single-pass extraction of all front matter:
        title, trans-title, authors (with ORCID/email/affiliation/country),
        DOI, received/accepted dates, abstract (ES + EN), keywords (ES + EN).
        """
        metadata: Dict = {
            'title': '',
            'trans_title': '',
            'authors': [],
            'doi': '',
            'received': None,
            'accepted': None,
            'abstract_es': '',
            'abstract_en': '',
            'keywords_es': [],
            'keywords_en': [],
            'volume': '',
            'issue': '',
            'elocation_id': '',
        }

        paragraphs = list(doc.paragraphs)

        # Locate where the body begins
        first_body_idx = len(paragraphs)
        for i, para in enumerate(paragraphs):
            if i < 3:
                continue
            style = getattr(getattr(para, 'style', None), 'name', '') or ''
            text_lower = (para.text or '').strip().lower()
            is_heading = style.startswith('Heading')
            is_section = text_lower in self._BODY_SECTION_LABELS and len(text_lower) < 100
            if (is_heading and is_section) or (is_section and not is_heading):
                first_body_idx = i
                break

        # State flags
        in_abstract_es = False
        in_abstract_en = False
        in_keywords_es = False
        in_keywords_en = False
        abstract_es_parts: List[str] = []
        abstract_en_parts: List[str] = []
        title_set = False
        check_trans_title = False

        current_author: Optional[Dict] = None

        kw_es_prefixes = ('palabras clave', 'palabras-clave', 'palabras claves')
        kw_en_prefixes = ('keywords', 'key words', 'keyword')

        for i, para in enumerate(paragraphs[:first_body_idx]):
            text = (para.text or '').strip()
            if not text:
                continue
            text_lower = text.lower()

            # ---- Title: first non-empty paragraph that is NOT article type
            if not title_set:
                if text_lower in self._ARTICLE_TYPES:
                    continue
                metadata['title'] = text
                title_set = True
                check_trans_title = True
                continue

            # ---- DOI ----------------------------------------------------
            doi_m = self._DOI_RE.search(text)
            if doi_m and not metadata['doi']:
                metadata['doi'] = doi_m.group(1).rstrip('.')

            # ---- Received date ------------------------------------------
            if any(w in text_lower for w in ('recibido', 'received', 'recepción')):
                d = self._parse_date(text)
                if d:
                    metadata['received'] = d

            # ---- Accepted date ------------------------------------------
            if any(w in text_lower for w in ('aceptado', 'accepted', 'aceptación')):
                d = self._parse_date(text)
                if d:
                    metadata['accepted'] = d

            # ---- Abstract Spanish header --------------------------------
            if text_lower in ('resumen', 'resumen:'):
                in_abstract_es = True
                in_abstract_en = in_keywords_es = in_keywords_en = False
                continue
            if text_lower.startswith('resumen') and ':' in text and not in_abstract_es:
                in_abstract_es = True
                in_abstract_en = in_keywords_es = in_keywords_en = False
                after = text.split(':', 1)[-1].strip()
                if after:
                    abstract_es_parts.append(after)
                continue

            # ---- Abstract English header --------------------------------
            if text_lower == 'abstract':
                in_abstract_en = True
                in_abstract_es = in_keywords_es = in_keywords_en = False
                continue
            if text_lower.startswith('abstract') and ':' in text and not in_abstract_en:
                in_abstract_en = True
                in_abstract_es = in_keywords_es = in_keywords_en = False
                after = text.split(':', 1)[-1].strip()
                if after:
                    abstract_en_parts.append(after)
                continue

            # ---- Keywords Spanish header --------------------------------
            if any(text_lower.startswith(p) for p in kw_es_prefixes):
                in_keywords_es = True
                in_keywords_en = in_abstract_es = in_abstract_en = False
                after = text.split(':', 1)[-1].strip() if ':' in text else ''
                for kw in re.split(r'[,;]', after):
                    kw = kw.strip().rstrip('.')
                    if kw and len(kw) < 80:
                        metadata['keywords_es'].append(kw)
                continue

            # ---- Keywords English header --------------------------------
            if any(text_lower.startswith(p) for p in kw_en_prefixes):
                in_keywords_en = True
                in_keywords_es = in_abstract_es = in_abstract_en = False
                after = text.split(':', 1)[-1].strip() if ':' in text else ''
                for kw in re.split(r'[,;]', after):
                    kw = kw.strip().rstrip('.')
                    if kw and len(kw) < 80:
                        metadata['keywords_en'].append(kw)
                continue

            # ---- Content inside abstract ES ----------------------------
            if in_abstract_es:
                abstract_es_parts.append(text)
                continue

            # ---- Content inside abstract EN ----------------------------
            if in_abstract_en:
                abstract_en_parts.append(text)
                continue

            # ---- Continuation of keywords ES ---------------------------
            if in_keywords_es:
                for kw in re.split(r'[,;]', text):
                    kw = kw.strip().rstrip('.')
                    if kw and len(kw) < 80:
                        metadata['keywords_es'].append(kw)
                continue

            # ---- Continuation of keywords EN ---------------------------
            if in_keywords_en:
                for kw in re.split(r'[,;]', text):
                    kw = kw.strip().rstrip('.')
                    if kw and len(kw) < 80:
                        metadata['keywords_en'].append(kw)
                continue

            # ---- ORCID line --------------------------------------------
            orcid_m = self._ORCID_RE.search(text)
            if orcid_m:
                orcid = orcid_m.group(1)
                target = current_author if current_author else (
                    metadata['authors'][-1] if metadata['authors'] else None
                )
                if target is not None:
                    target['orcid'] = orcid
                # Also capture country from the same line (e.g. "Perú https://orcid.org/...")
                leftover = self._URL_RE.sub('', self._ORCID_RE.sub('', text)).strip()
                if leftover and leftover.lower() in self._COUNTRY_NAMES:
                    if target is not None:
                        target.setdefault('country', leftover)
                continue

            # ---- Email line --------------------------------------------
            email_m = self._EMAIL_RE.search(text)
            if email_m and text.count(' ') < 4:
                target = current_author if current_author else (
                    metadata['authors'][-1] if metadata['authors'] else None
                )
                if target is not None:
                    target['email'] = email_m.group(0)
                continue

            # ---- Lines to skip (URL-only, date markers) ----------------
            if text.startswith(('http://', 'https://', 'www.', 'URL:', 'DOI:', '©')):
                continue
            recv_keywords = ('recibido', 'recepción', 'received')
            acc_keywords = ('aceptado', 'aceptación', 'accepted')
            is_date_line = (
                (
                    (bool(self._DATE_ES_RE.search(text)) or bool(self._DATE_SLASH_RE.search(text)))
                    and any(w in text_lower for w in recv_keywords + acc_keywords)
                )
                or text_lower.startswith((
                    'recibido:', 'aceptado:', 'received:', 'accepted:',
                    'recibido ', 'aceptado ',
                ))
                or (
                    any(w in text_lower for w in recv_keywords)
                    and any(w in text_lower for w in acc_keywords)
                )
            )
            if is_date_line:
                continue

            # ---- Trans-title: second non-empty paragraph candidate ------
            if check_trans_title and len(text) > 15:
                check_trans_title = False
                has_url = text.startswith('http')
                has_orcid = bool(self._ORCID_RE.search(text))
                has_date = bool(self._DATE_ES_RE.search(text)) or bool(self._DATE_SLASH_RE.search(text))
                is_inst = any(w in text_lower for w in self._INST_KEYWORDS)
                if not (has_url or has_orcid or has_date or is_inst):
                    metadata['trans_title'] = text
                    continue

            # ---- Author / affiliation / country lines ------------------
            if not re.match(r'^[A-ZÁÉÍÓÚÑ]', text) or len(text) > 250:
                continue

            is_institution = any(w in text_lower for w in self._INST_KEYWORDS)
            is_country = text_lower.strip() in self._COUNTRY_NAMES or (
                len(text.split()) <= 3 and text_lower.strip() in self._COUNTRY_NAMES
            )

            if is_institution:
                # Try to split "Author Name Universidad XYZ" into name + aff
                if current_author is None and not metadata['authors']:
                    for kw in self._INST_KEYWORDS:
                        idx = text_lower.find(kw)
                        if idx > 5:
                            name_part = text[:idx].strip().rstrip(',').strip()
                            aff_part = text[idx:].strip()
                            if name_part:
                                current_author = {
                                    'name_raw': name_part,
                                    'orcid': '', 'email': '',
                                    'aff': aff_part, 'country': '',
                                }
                                break
                    else:
                        # Couldn't split – store as aff of a generic author
                        if not metadata['authors']:
                            current_author = {
                                'name_raw': '', 'orcid': '', 'email': '',
                                'aff': text, 'country': '',
                            }
                else:
                    target = current_author if current_author else (
                        metadata['authors'][-1] if metadata['authors'] else None
                    )
                    if target is not None:
                        target['aff'] = text
            elif is_country:
                target = current_author if current_author else (
                    metadata['authors'][-1] if metadata['authors'] else None
                )
                if target is not None:
                    target['country'] = text
            else:
                # Treat as a new author name
                if current_author:
                    metadata['authors'].append(current_author)
                current_author = {
                    'name_raw': text,
                    'orcid': '', 'email': '', 'aff': '', 'country': '',
                }

        # Flush last author
        if current_author:
            metadata['authors'].append(current_author)

        metadata['abstract_es'] = '\n\n'.join(abstract_es_parts)
        metadata['abstract_en'] = '\n\n'.join(abstract_en_parts)

        if not metadata['title'] and input_path:
            metadata['title'] = Path(input_path).stem or 'Sin título'
        if not metadata['title']:
            metadata['title'] = 'Sin título'

        return metadata

    # ------------------------------------------------------------------ #
    # Body extraction
    # ------------------------------------------------------------------ #

    def _extract_body(
        self,
        doc: Document,
    ) -> List[Tuple[str, List[Union[str, List[List[str]], Dict]]]]:
        """
        Extract body sections with embedded images and tables at their physical positions.
        Sequential numbering: figures → fig-1.jpg, fig-2.jpg…; tables → table-1.jpg, table-2.jpg…
        Captions adjacent to media items are associated automatically.
        """
        sections: List[Tuple[str, List[Union[str, List[List[str]], Dict]]]] = []
        current_section: Optional[str] = None
        current_items: List[Union[str, List[List[str]], Dict]] = []
        pre_body = True
        fig_seq = 0
        table_seq = 0

        def flush() -> None:
            if current_section is not None and current_items:
                sections.append((current_section, list(current_items)))

        all_blocks = list(doc.iter_inner_content())
        consumed: set = set()

        for idx, block in enumerate(all_blocks):
            if idx in consumed:
                continue

            # ---- TABLE -------------------------------------------------------
            if isinstance(block, Table):
                if not pre_body and current_section is not None:
                    table_seq += 1
                    seq = str(table_seq)
                    caption_text = ''
                    caption_num = None
                    # Look ahead: next paragraph may be the caption
                    if idx + 1 < len(all_blocks) and not isinstance(all_blocks[idx + 1], Table):
                        nxt = (all_blocks[idx + 1].text or '').strip()
                        tm = self._TABLE_CAPTION_RE.match(nxt)
                        if tm:
                            caption_num = tm.group(2).lower()
                            caption_text = nxt
                            consumed.add(idx + 1)
                    current_items.append({
                        'kind': 'table_fig',
                        'label': f'Tabla {seq}',
                        'caption': caption_text,
                        'rid': f't{seq}',
                        'filename': f'table-{seq}.jpg',
                        'caption_num': caption_num,
                    })
                continue

            # ---- PARAGRAPH ---------------------------------------------------
            text = (block.text or '').strip()
            style = getattr(getattr(block, 'style', None), 'name', '') or ''
            text_lower = text.lower()
            is_heading = style.startswith('Heading')

            # Stop at references section
            if any(text_lower == lbl or text_lower.startswith(lbl) for lbl in self._REFS_LABELS):
                break

            # Section heading → flush + start new section
            is_body_section = text_lower in self._BODY_SECTION_LABELS and len(text) < 100
            if is_heading or is_body_section:
                if pre_body and is_body_section:
                    pre_body = False
                if not pre_body:
                    flush()
                    current_section = text or 'Sección'
                    current_items = []
                continue

            if pre_body or current_section is None:
                continue

            # ---- Inline image ------------------------------------------------
            if self._has_inline_image(block):
                fig_seq += 1
                seq = str(fig_seq)
                caption_text = ''
                caption_num = None
                # Case 1: caption text in the same paragraph as the image
                fig_m = self._FIG_CAPTION_RE.match(text)
                if fig_m:
                    caption_num = fig_m.group(2).lower()
                    caption_text = text
                else:
                    # Case 2: caption in the next paragraph
                    if idx + 1 < len(all_blocks) and not isinstance(all_blocks[idx + 1], Table):
                        nxt = (all_blocks[idx + 1].text or '').strip()
                        fm = self._FIG_CAPTION_RE.match(nxt)
                        if fm:
                            caption_num = fm.group(2).lower()
                            caption_text = nxt
                            consumed.add(idx + 1)
                current_items.append({
                    'kind': 'fig',
                    'label': f'Figura {seq}',
                    'caption': caption_text,
                    'rid': f'f{seq}',
                    'filename': f'fig-{seq}.jpg',
                    'caption_num': caption_num,
                })
                continue

            # ---- Regular text (skip pure caption paragraphs) -----------------
            # A standalone caption para (no adjacent image detected) is kept as
            # body text so its "Figura N" label remains clickable via xref.
            if text:
                current_items.append(text)

        flush()

        if not sections:
            sections = self._extract_body_fallback(doc)

        return sections

    def _extract_body_fallback(
        self, doc: Document
    ) -> List[Tuple[str, List[Union[str, List[List[str]], Dict]]]]:
        """Fallback body extraction with inline image and table detection."""
        sections: List[Tuple[str, List[Union[str, List[List[str]], Dict]]]] = []
        current_section: Optional[str] = None
        current_items: List[Union[str, List[List[str]], Dict]] = []
        fig_seq = 0
        table_seq = 0

        def flush() -> None:
            if current_section is not None and current_items:
                sections.append((current_section, list(current_items)))

        all_blocks = list(doc.iter_inner_content())
        consumed: set = set()

        for idx, block in enumerate(all_blocks):
            if idx in consumed:
                continue

            if isinstance(block, Table):
                if current_section is not None:
                    table_seq += 1
                    seq = str(table_seq)
                    caption_text = ''
                    caption_num = None
                    if idx + 1 < len(all_blocks) and not isinstance(all_blocks[idx + 1], Table):
                        nxt = (all_blocks[idx + 1].text or '').strip()
                        tm = self._TABLE_CAPTION_RE.match(nxt)
                        if tm:
                            caption_num = tm.group(2).lower()
                            caption_text = nxt
                            consumed.add(idx + 1)
                    current_items.append({
                        'kind': 'table_fig',
                        'label': f'Tabla {seq}',
                        'caption': caption_text,
                        'rid': f't{seq}',
                        'filename': f'table-{seq}.jpg',
                        'caption_num': caption_num,
                    })
                continue

            text = (block.text or '').strip()
            style = getattr(getattr(block, 'style', None), 'name', '') or ''
            text_lower = text.lower()

            if style.startswith('Heading') or text_lower in self._BODY_SECTION_LABELS:
                flush()
                current_section = text or 'Sección'
                current_items = []
                continue

            if current_section is None:
                continue

            if self._has_inline_image(block):
                fig_seq += 1
                seq = str(fig_seq)
                caption_text = ''
                caption_num = None
                fig_m = self._FIG_CAPTION_RE.match(text)
                if fig_m:
                    caption_num = fig_m.group(2).lower()
                    caption_text = text
                else:
                    if idx + 1 < len(all_blocks) and not isinstance(all_blocks[idx + 1], Table):
                        nxt = (all_blocks[idx + 1].text or '').strip()
                        fm = self._FIG_CAPTION_RE.match(nxt)
                        if fm:
                            caption_num = fm.group(2).lower()
                            caption_text = nxt
                            consumed.add(idx + 1)
                current_items.append({
                    'kind': 'fig',
                    'label': f'Figura {seq}',
                    'caption': caption_text,
                    'rid': f'f{seq}',
                    'filename': f'fig-{seq}.jpg',
                    'caption_num': caption_num,
                })
                continue

            if text:
                current_items.append(text)

        flush()
        return sections

    def _media_number_sort_key(self, token: str) -> Tuple[int, str]:
        """Sort media ids naturally: 2 < 10 and 2a after 2."""
        m = re.match(r'^(\d+)([a-z]?)$', token or '', re.IGNORECASE)
        if not m:
            return (10**9, token or '')
        return (int(m.group(1)), (m.group(2) or '').lower())

    def _media_kind_from_prefix(self, prefix: str) -> str:
        """Resolve textual media prefix to internal kind."""
        normalized = self._strip_accents((prefix or '').lower()).replace('.', '')
        return 'table_fig' if normalized.startswith('tab') else 'fig'

    def _expand_media_id_sequence(self, text: str) -> List[str]:
        """
        Expand a sequence like '1, 2 y 3' or '1-3' into media ids.
        Supports numeric ids with optional trailing letter (e.g., 2a).
        """
        cleaned = (text or '').lower().strip()
        if not cleaned:
            return []
        cleaned = cleaned.replace('–', '-')
        cleaned = re.sub(r'\b(?:y|e|and|&)\b', ',', cleaned)
        cleaned = re.sub(r'\s+', ' ', cleaned)

        ids: List[str] = []
        for part in (p.strip() for p in cleaned.split(',')):
            if not part:
                continue
            token = re.sub(r'^(?:n(?:[°º]|o\.?))\s*', '', part).strip()
            if not token:
                continue

            range_match = re.match(r'^(\d+)-(\d+)$', token)
            if range_match:
                start = int(range_match.group(1))
                end = int(range_match.group(2))
                if 0 < start <= end and (end - start) <= 50:
                    for n in range(start, end + 1):
                        ids.append(str(n))
                continue

            token_match = re.match(r'^(\d+[a-z]?)$', token)
            if token_match:
                ids.append(token_match.group(1))

        # Preserve order and deduplicate
        seen = set()
        deduped: List[str] = []
        for token in ids:
            if token not in seen:
                seen.add(token)
                deduped.append(token)
        return deduped

    def _extract_media_references_from_text(self, text: str) -> List[Tuple[str, str]]:
        """Extract media references from text as (kind, id_token)."""
        refs: List[Tuple[str, str]] = []
        if not text:
            return refs

        for m in self._MEDIA_REF_BLOCK_RE.finditer(text):
            kind = self._media_kind_from_prefix(m.group(1))
            seq = m.group(2) or ''
            for token in self._expand_media_id_sequence(seq):
                refs.append((kind, token))
        return refs

    def _has_inline_image(self, para) -> bool:
        """
        Return True only if the paragraph contains an actual embedded image.
        Uses a:blip (DrawingML) or v:imagedata (VML legacy) as the signal.
        These are present exclusively for raster/vector images (PNG, JPG, WMF, EMF)
        and are absent for charts, SmartArt, text boxes, WordArt, etc.
        """
        el = para._element
        # DrawingML images: a:blip holds the image relationship
        NS_A = 'http://schemas.openxmlformats.org/drawingml/2006/main'
        if el.findall(f'.//{{{NS_A}}}blip'):
            return True
        # VML legacy images: v:imagedata
        NS_V = 'urn:schemas-microsoft-com:vml'
        if el.findall(f'.//{{{NS_V}}}imagedata'):
            return True
        return False

    def _inject_external_media_references(
        self,
        doc: Document,
        body_sections: List[Tuple[str, List[Union[str, List[List[str]], Dict]]]],
    ) -> None:
        """
        Add placeholder media nodes for text references NOT already detected as
        physical embedded images/tables by _extract_body.

        Physical items (detected inline) take priority; this method only adds
        items for references that have NO corresponding physical element.
        Naming convention: Figura N → fig-N.jpg, Tabla N → table-N.jpg.
        """
        if not body_sections:
            body_sections.append(('Contenido', []))

        # Collect rids already present from physical detection
        physical_fig_rids: set = set()
        physical_table_rids: set = set()
        # Also collect caption_nums already covered
        covered_fig_nums: set = set()
        covered_table_nums: set = set()

        for _sec, items in body_sections:
            for item in items:
                if not isinstance(item, dict):
                    continue
                rid = item.get('rid', '')
                cap_num = (item.get('caption_num') or '').lower().strip()
                if item.get('kind') == 'fig':
                    physical_fig_rids.add(rid)
                    # Sequential label number
                    m = re.match(r'^f(\d+[a-z]?)$', rid)
                    if m:
                        covered_fig_nums.add(m.group(1))
                    if cap_num:
                        covered_fig_nums.add(cap_num)
                elif item.get('kind') == 'table_fig':
                    physical_table_rids.add(rid)
                    m = re.match(r'^t(\d+[a-z]?)$', rid)
                    if m:
                        covered_table_nums.add(m.group(1))
                    if cap_num:
                        covered_table_nums.add(cap_num)

        # Determine next sequential numbers beyond physical items
        next_fig = len(physical_fig_rids) + 1
        next_table = len(physical_table_rids) + 1

        refs_by_kind: Dict[str, set] = {'fig': set(), 'table_fig': set()}

        # Scan full document text for Figura/Tabla references
        for para in doc.paragraphs:
            text = (para.text or '').strip()
            if not text:
                continue
            for kind, number in self._extract_media_references_from_text(text):
                refs_by_kind[kind].add(number)

        # Also scan extracted body text items
        for _section_title, items in body_sections:
            for item in items:
                if not isinstance(item, str):
                    continue
                for kind, number in self._extract_media_references_from_text(item):
                    refs_by_kind[kind].add(number)

        # Add only items NOT already covered by physical detection
        media_items: List[Dict[str, str]] = []
        for number in sorted(refs_by_kind['fig'], key=self._media_number_sort_key):
            if number in covered_fig_nums:
                continue
            seq = str(next_fig)
            next_fig += 1
            media_items.append({
                'kind': 'fig',
                'label': f'Figura {seq}',
                'caption': '',
                'rid': f'f{seq}',
                'filename': f'fig-{seq}.jpg',
                'caption_num': number,
            })
        for number in sorted(refs_by_kind['table_fig'], key=self._media_number_sort_key):
            if number in covered_table_nums:
                continue
            seq = str(next_table)
            next_table += 1
            media_items.append({
                'kind': 'table_fig',
                'label': f'Tabla {seq}',
                'caption': '',
                'rid': f't{seq}',
                'filename': f'table-{seq}.jpg',
                'caption_num': number,
            })

        if media_items:
            body_sections[-1][1].extend(media_items)

    # ------------------------------------------------------------------ #
    # References extraction
    # ------------------------------------------------------------------ #

    def _extract_references(self, doc: Document) -> List[str]:
        """Extract, deduplicate and consolidate bibliography entries."""
        raw: List[str] = []
        in_refs = False

        for para in doc.paragraphs:
            text = (para.text or '').strip()
            text_lower = text.lower()

            if any(
                text_lower == lbl or text_lower.startswith(lbl)
                for lbl in ('referencias', 'bibliografía', 'bibliography', 'references',
                            'fuentes bibliográficas')
            ):
                in_refs = True
                continue

            if in_refs and text:
                raw.append(self._deduplicate_text(text))

        return self._consolidate_references(raw)

    # ------------------------------------------------------------------ #
    # Reference citation builders
    # ------------------------------------------------------------------ #

    def _build_mixed_citation_with_links(self, parent: etree.Element, ref_text: str) -> None:
        """Add <mixed-citation> with URLs wrapped in <ext-link>."""
        if not ref_text or not ref_text.strip():
            return
        url_re = re.compile(r'https?://[^\s<>"\']+')
        parts: List[Tuple[str, str]] = []
        last = 0
        for m in url_re.finditer(ref_text):
            if m.start() > last:
                parts.append(('text', ref_text[last:m.start()]))
            parts.append(('url', m.group(0)))
            last = m.end()
        if last < len(ref_text):
            parts.append(('text', ref_text[last:]))

        mc = etree.SubElement(parent, 'mixed-citation')
        if not parts:
            mc.text = ref_text
            return

        pending = ''
        for idx, (kind, content) in enumerate(parts):
            if kind == 'text':
                pending = content
                if idx == 0:
                    mc.text = pending
            else:
                if idx == 0:
                    mc.text = ''
                el = etree.SubElement(mc, 'ext-link', {
                    'ext-link-type': 'uri',
                    f'{{{self.XLINK_NS}}}href': content,
                })
                el.text = content
                el.tail = pending
                pending = ''
        if pending:
            if len(mc):
                mc[-1].tail = (mc[-1].tail or '') + pending
            else:
                mc.text = (mc.text or '') + pending

    def _parse_reference_structured(self, ref_text: str) -> Optional[Dict]:
        """Parse a reference string into structured fields for <element-citation>."""
        year_m = self._YEAR_RE.search(ref_text)
        if not year_m:
            return None
        year = year_m.group(0)

        doi = ''
        doi_m = self._DOI_RE.search(ref_text)
        if doi_m:
            doi = doi_m.group(1).rstrip('.')

        url = ''
        url_m = self._URL_RE.search(ref_text)
        if url_m:
            url = url_m.group(0)

        # Authors text: everything before the year (and opening paren)
        authors_text = ref_text[:year_m.start()].rstrip('(').strip()

        # After the year
        after_year = ref_text[year_m.end():].lstrip(')').lstrip('.').strip()
        chunks = re.split(r'\.\s+', after_year, maxsplit=2)
        title = chunks[0].strip() if chunks else ''
        source = chunks[1].strip() if len(chunks) > 1 else ''

        # Volume/issue
        vol_m = re.search(r',\s*(\d+)\s*[\(\[](\d+)[\)\]]', after_year)
        volume = vol_m.group(1) if vol_m else ''
        issue = vol_m.group(2) if vol_m else ''

        # Pages
        pages_m = re.search(r'(?:pp?\.\s*)?(\d+)[–\-](\d+)', after_year)
        fpage = pages_m.group(1) if pages_m else ''
        lpage = pages_m.group(2) if pages_m else ''

        # Parse individual author names
        author_list: List[Dict[str, str]] = []
        for part in re.split(r';\s*|(?:,\s*)?(?:&|y|and)\s+(?=[A-ZÁÉÍÓÚÑ])', authors_text):
            part = part.strip().rstrip(',')
            if not part or len(part) < 2:
                continue
            comma = part.split(',', 1)
            if len(comma) == 2 and len(comma[0].split()) <= 3:
                author_list.append({'surname': comma[0].strip(), 'given': comma[1].strip()})
            else:
                words = part.split()
                if len(words) >= 2:
                    author_list.append({'surname': words[-1], 'given': ' '.join(words[:-1])})
                elif words:
                    author_list.append({'surname': part, 'given': ''})

        return {
            'authors': author_list,
            'year': year,
            'title': title,
            'source': source,
            'volume': volume,
            'issue': issue,
            'fpage': fpage,
            'lpage': lpage,
            'doi': doi,
            'url': url,
        }

    def _build_element_citation(self, parent: etree.Element, ref_text: str) -> None:
        """Add <element-citation> with structured metadata when parseable."""
        parsed = self._parse_reference_structured(ref_text)
        if not parsed:
            return

        # Determine publication type
        pub_type = 'journal'
        if any(w in ref_text.lower() for w in ('tesis', 'thesis', 'maestría', 'master', 'doctoral')):
            pub_type = 'thesis'
        elif not parsed['volume'] and not parsed['issue'] and parsed['source']:
            pub_type = 'book'

        ec = etree.SubElement(parent, 'element-citation', {'publication-type': pub_type})

        for auth in parsed['authors']:
            pg = etree.SubElement(ec, 'person-group', {'person-group-type': 'author'})
            name_el = etree.SubElement(pg, 'name')
            sn = etree.SubElement(name_el, 'surname')
            sn.text = auth['surname']
            if auth['given']:
                gn = etree.SubElement(name_el, 'given-names')
                gn.text = auth['given']

        if parsed['year']:
            yr = etree.SubElement(ec, 'year')
            yr.text = parsed['year']
        if parsed['title']:
            at = etree.SubElement(ec, 'article-title')
            at.text = parsed['title']
        if parsed['source']:
            src = etree.SubElement(ec, 'source')
            src.text = parsed['source']
        if parsed['volume']:
            v = etree.SubElement(ec, 'volume')
            v.text = parsed['volume']
        if parsed['issue']:
            iss = etree.SubElement(ec, 'issue')
            iss.text = parsed['issue']
        if parsed['fpage']:
            fp = etree.SubElement(ec, 'fpage')
            fp.text = parsed['fpage']
        if parsed['lpage']:
            lp = etree.SubElement(ec, 'lpage')
            lp.text = parsed['lpage']

        link_url = parsed['url'] or (f'https://doi.org/{parsed["doi"]}' if parsed['doi'] else '')
        if link_url:
            el = etree.SubElement(ec, 'ext-link', {
                'ext-link-type': 'uri',
                f'{{{self.XLINK_NS}}}href': link_url,
            })
            el.text = link_url

    # ------------------------------------------------------------------ #
    # In-text citation helpers
    # ------------------------------------------------------------------ #

    def _strip_accents(self, s: str) -> str:
        """Remove diacritics for accent-insensitive surname matching."""
        return ''.join(
            c for c in unicodedata.normalize('NFD', s)
            if unicodedata.category(c) != 'Mn'
        )

    def _build_citation_index(self, references: List[str]) -> Dict[Tuple[str, str], str]:
        """
        Build {(normalized_first_surname, year): 'B{n}'} from the reference list.
        Used to resolve in-text citations to their <ref id> in the back matter.
        """
        index: Dict[Tuple[str, str], str] = {}
        for i, ref in enumerate(references, 1):
            year_m = self._YEAR_RE.search(ref)
            if not year_m:
                continue
            year = year_m.group(0)
            before = ref[:year_m.start()]
            name_m = re.search(r'([A-ZÁÉÍÓÚÑ][a-záéíóúñ]+)', before)
            if name_m:
                key = (self._strip_accents(name_m.group(1).lower()), year)
                if key not in index:          # keep first occurrence
                    index[key] = f'B{i}'
        return index

    def _lookup_citation(
        self, surname: str, year: str, index: Dict[Tuple[str, str], str]
    ) -> Optional[str]:
        """Find a reference id by surname + year, with accent-insensitive fallback."""
        norm = self._strip_accents(surname.lower())
        rid = index.get((norm, year))
        if rid:
            return rid
        # Try only the first word of a compound surname
        first = norm.split()[0] if ' ' in norm else norm
        return index.get((first, year))

    def _build_media_index(
        self,
        body_sections: List[Tuple[str, List[Union[str, List[List[str]], Dict]]]],
    ) -> Dict[str, Dict[str, str]]:
        """
        Build label->(rid, ref-type) map for figure/table references.
        Keys are normalized variants: "figura 1", "fig 1", "tabla 2a", "tab 2a".
        """
        media_index: Dict[str, Dict[str, str]] = {}

        def norm(s: str) -> str:
            s = self._strip_accents((s or '').lower())
            s = re.sub(r'\s+', ' ', s).strip()
            return s

        for _section_title, items in body_sections:
            for item in items:
                if not isinstance(item, dict):
                    continue
                rid = (item.get('rid') or '').strip()
                label = (item.get('label') or '').strip()
                if not rid or not label:
                    continue
                m = re.match(r'^(figura|fig\.?|tabla|tab\.?)\s+(\d+[a-z]?)$', label, re.IGNORECASE)
                if not m:
                    continue
                prefix = m.group(1).lower().replace('.', '')
                num = m.group(2).lower()
                # caption_num: the original number from the Word caption (may differ from seq num)
                cap_num = (item.get('caption_num') or '').lower().strip()
                if prefix in ('figura', 'fig'):
                    media_index[norm(f'figura {num}')] = {'rid': rid, 'ref_type': 'fig'}
                    media_index[norm(f'fig {num}')] = {'rid': rid, 'ref_type': 'fig'}
                    # Also register the original caption number so "ver Figura 3" still resolves
                    # when the sequential label is "Figura 2" (i.e., it's the 2nd physical image)
                    if cap_num and cap_num != num:
                        media_index[norm(f'figura {cap_num}')] = {'rid': rid, 'ref_type': 'fig'}
                        media_index[norm(f'fig {cap_num}')] = {'rid': rid, 'ref_type': 'fig'}
                else:
                    # Tables are managed as figure-like dependent images for OJS.
                    media_index[norm(f'tabla {num}')] = {'rid': rid, 'ref_type': 'fig'}
                    media_index[norm(f'tab {num}')] = {'rid': rid, 'ref_type': 'fig'}
                    if cap_num and cap_num != num:
                        media_index[norm(f'tabla {cap_num}')] = {'rid': rid, 'ref_type': 'fig'}
                        media_index[norm(f'tab {cap_num}')] = {'rid': rid, 'ref_type': 'fig'}
        return media_index

    def _add_para_with_xrefs(
        self,
        parent_el: etree.Element,
        text: str,
        citation_index: Dict[Tuple[str, str], str],
        media_index: Optional[Dict[str, Dict[str, str]]] = None,
    ) -> None:
        """
        Create a <p> where bibliography and figure/table references become clickable:
        - <xref ref-type="bibr" ...> for author-year citations
        - <xref ref-type="fig" ...>  for Figura/Tabla references
        """
        p = etree.SubElement(parent_el, 'p')
        if not text:
            p.text = text or ''
            return
        media_index = media_index or {}

        matches: List[Tuple[int, int, str, re.Match]] = []
        for m in self._CITATION_RE.finditer(text):
            matches.append((m.start(), m.end(), 'bibr', m))
        for m in self._FIGREF_RE.finditer(text):
            matches.append((m.start(), m.end(), 'fig', m))
        matches.sort(key=lambda x: (x[0], -(x[1] - x[0])))

        # Remove overlaps by keeping the first match by sorted order
        clean: List[Tuple[int, int, str, re.Match]] = []
        cursor = -1
        for start, end, kind, m in matches:
            if start < cursor:
                continue
            clean.append((start, end, kind, m))
            cursor = end

        if not clean:
            p.text = text
            return

        last_end = 0
        last_el: Optional[etree.Element] = None
        def append_text(s: str) -> None:
            if not s:
                return
            if last_el is None:
                p.text = (p.text or '') + s
            else:
                last_el.tail = (last_el.tail or '') + s

        for start, end, kind, m in clean:
            append_text(text[last_end:start])

            if kind == 'fig':
                ref_label = m.group(1)
                ref_num = m.group(2).lower()
                prefix = self._strip_accents(ref_label.lower().replace('.', ''))
                key = f'figura {ref_num}' if prefix in ('figura', 'fig') else f'tabla {ref_num}'
                media_ref = media_index.get(key)
                if media_ref:
                    xref = etree.SubElement(
                        p,
                        'xref',
                        {'ref-type': media_ref['ref_type'], 'rid': media_ref['rid']},
                    )
                    xref.text = m.group(0)
                    last_el = xref
                else:
                    append_text(m.group(0))
            elif citation_index:
                if m.group(1):
                    # Pattern A: Author (YYYY) -> xref wraps year only
                    author_text = m.group(1)
                    year = m.group(2)
                    first_surname = re.search(r'[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+', author_text)
                    rid = (
                        self._lookup_citation(first_surname.group(0), year, citation_index)
                        if first_surname else None
                    )
                    if rid:
                        full = m.group(0)
                        paren_idx = full.index('(')
                        pre_paren = full[:paren_idx]
                        inside = full[paren_idx + 1:-1]
                        year_pos = inside.index(year)
                        after_year = inside[year_pos + len(year):]
                        append_text(pre_paren + '(')
                        xref = etree.SubElement(p, 'xref', {'ref-type': 'bibr', 'rid': rid})
                        xref.text = year
                        xref.tail = after_year + ')'
                        last_el = xref
                    else:
                        append_text(m.group(0))
                else:
                    # Pattern B: (Author..., YYYY) -> xref wraps inner text
                    author_text = m.group(3)
                    year = m.group(4)
                    first_surname = re.search(r'[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+', author_text)
                    rid = (
                        self._lookup_citation(first_surname.group(0), year, citation_index)
                        if first_surname else None
                    )
                    if rid:
                        inner = m.group(0)[1:-1]
                        append_text('(')
                        xref = etree.SubElement(p, 'xref', {'ref-type': 'bibr', 'rid': rid})
                        xref.text = inner
                        xref.tail = ')'
                        last_el = xref
                    else:
                        append_text(m.group(0))
            else:
                append_text(m.group(0))

            last_end = end

        append_text(text[last_end:])

    # ------------------------------------------------------------------ #
    # JATS XML builder
    # ------------------------------------------------------------------ #

    def _build_jats_xml(
        self,
        metadata: Dict,
        body_sections: List[Tuple[str, List[Union[str, List[List[str]], Dict]]]],
        references: List[str],
        journal_meta: Optional[Dict] = None,
    ) -> etree.Element:
        """Build a complete JATS 1.1 XML tree from extracted content."""
        fig_count = sum(
            1
            for _s, items in body_sections
            for item in items
            if isinstance(item, dict) and item.get('kind') in ('fig', 'table_fig')
        )
        # In this workflow, tables are treated as image-like figures.
        table_count = 0

        title_text = (metadata.get('title') or 'Sin título').strip() or 'Sin título'
        trans_title = (metadata.get('trans_title') or '').strip()
        authors = metadata.get('authors') or []
        doi = (metadata.get('doi') or '').strip()
        received = metadata.get('received')
        accepted = metadata.get('accepted')
        abstract_es = (metadata.get('abstract_es') or '').strip()
        abstract_en = (metadata.get('abstract_en') or '').strip()
        keywords_es = metadata.get('keywords_es') or []
        keywords_en = metadata.get('keywords_en') or []
        volume = (metadata.get('volume') or '').strip()
        issue = (metadata.get('issue') or '').strip()
        elocation_id = (metadata.get('elocation_id') or '').strip()

        # ---- Root -------------------------------------------------------
        article = etree.Element('article', {
            'article-type': 'research-article',
            'dtd-version': '1.1',
            'specific-use': 'sps-1.9',
            f'{{{self.XML_NS}}}lang': 'es',
        }, nsmap={
            'xlink': self.XLINK_NS,
            'mml': 'http://www.w3.org/1998/Math/MathML',
        })

        # ---- Front ------------------------------------------------------
        front = etree.SubElement(article, 'front')

        # journal-meta: journal_meta (multi-revista) > config (revista única) > placeholders OJS
        # Los placeholders {$journalTitle} y {$siteTitle} son reemplazados automáticamente
        # por el LensGalleyPlugin de OJS al servir el XML.
        if journal_meta:
            journal_id = str(journal_meta.get('journal_id') or '').strip() or '{$journalTitle}'
            journal_title = str(journal_meta.get('journal_title') or '').strip() or '{$journalTitle}'
            abbrev_title = str(journal_meta.get('journal_abbrev') or journal_title).strip() or '{$journalTitle}'
            publisher_name = str(journal_meta.get('publisher_name') or '').strip() or '{$siteTitle}'
            issn = str(journal_meta.get('issn') or '').strip()
        else:
            journal_id = (getattr(settings, 'JATS_JOURNAL_ID', '') or '').strip() or '{$journalTitle}'
            journal_title = (getattr(settings, 'JATS_JOURNAL_TITLE', '') or '').strip() or '{$journalTitle}'
            abbrev_title = journal_title
            publisher_name = (getattr(settings, 'JATS_PUBLISHER_NAME', '') or '').strip() or '{$siteTitle}'
            issn = (getattr(settings, 'JATS_ISSN', '') or '').strip()

        jm = etree.SubElement(front, 'journal-meta')
        ji = etree.SubElement(jm, 'journal-id', {'journal-id-type': 'publisher-id'})
        ji.text = journal_id
        jtg = etree.SubElement(jm, 'journal-title-group')
        jt = etree.SubElement(jtg, 'journal-title')
        jt.text = journal_title
        ab = etree.SubElement(jtg, 'abbrev-journal-title', {'abbrev-type': 'publisher'})
        ab.text = abbrev_title
        if issn:
            issn_el = etree.SubElement(jm, 'issn', {'pub-type': 'epub'})
            issn_el.text = issn
        pub = etree.SubElement(jm, 'publisher')
        pn = etree.SubElement(pub, 'publisher-name')
        pn.text = publisher_name

        # article-meta
        am = etree.SubElement(front, 'article-meta')

        # article-id
        if doi:
            aid = etree.SubElement(am, 'article-id', {'pub-id-type': 'doi'})
            aid.text = doi
        else:
            aid = etree.SubElement(am, 'article-id', {'pub-id-type': 'other'})
            aid.text = 'converted-1'

        # article-categories
        ac = etree.SubElement(am, 'article-categories')
        sg = etree.SubElement(ac, 'subj-group', {'subj-group-type': 'heading'})
        subj = etree.SubElement(sg, 'subject')
        subj.text = 'Artículo original'

        # title-group
        tg = etree.SubElement(am, 'title-group')
        at_el = etree.SubElement(tg, 'article-title')
        at_el.text = title_text
        if trans_title:
            ttg = etree.SubElement(tg, 'trans-title-group', {f'{{{self.XML_NS}}}lang': 'en'})
            tt = etree.SubElement(ttg, 'trans-title')
            tt.text = trans_title

        # contrib-group
        cg = etree.SubElement(am, 'contrib-group')
        if not authors:
            authors = [{'name_raw': 'Autor', 'orcid': '', 'email': '', 'aff': '', 'country': ''}]

        for idx, author in enumerate(authors):
            contrib = etree.SubElement(cg, 'contrib', {'contrib-type': 'author'})

            orcid = (author.get('orcid') or '').strip()
            if orcid:
                cid = etree.SubElement(contrib, 'contrib-id', {'contrib-id-type': 'orcid'})
                cid.text = orcid

            surname, given_names = self._parse_author_name(author.get('name_raw') or '')
            name_el = etree.SubElement(contrib, 'name')
            sn_el = etree.SubElement(name_el, 'surname')
            sn_el.text = surname
            if given_names:
                gn_el = etree.SubElement(name_el, 'given-names')
                gn_el.text = given_names

            xref_aff = etree.SubElement(contrib, 'xref', {'ref-type': 'aff', 'rid': f'aff{idx+1}'})
            sup_a = etree.SubElement(xref_aff, 'sup')
            sup_a.text = str(idx + 1)

            if idx == 0:
                xref_c = etree.SubElement(contrib, 'xref', {'ref-type': 'corresp', 'rid': 'c1'})
                sup_c = etree.SubElement(xref_c, 'sup')
                sup_c.text = '*'

        # Affiliations (inside contrib-group, after all contribs)
        for idx, author in enumerate(authors):
            aff_el = etree.SubElement(cg, 'aff', {'id': f'aff{idx+1}'})
            lbl = etree.SubElement(aff_el, 'label')
            lbl.text = str(idx + 1)

            aff_text = (author.get('aff') or '').strip()
            country_text = (author.get('country') or '').strip()

            inst_orig = etree.SubElement(aff_el, 'institution', {'content-type': 'original'})
            inst_orig.text = aff_text if aff_text else 'Sin especificar'

            if aff_text:
                org_name = aff_text.split('.')[0].strip() if '.' in aff_text else aff_text
                inst_org = etree.SubElement(aff_el, 'institution', {'content-type': 'orgname'})
                inst_org.text = org_name

            if country_text:
                country_el = etree.SubElement(aff_el, 'country')
                country_el.text = country_text

        # author-notes
        an = etree.SubElement(am, 'author-notes')
        corresp_el = etree.SubElement(an, 'corresp', {'id': 'c1'})
        lbl_c = etree.SubElement(corresp_el, 'label')
        lbl_c.text = '*'
        lbl_c.tail = ' Autor para la correspondencia: '
        first_email = next((a.get('email') for a in authors if a.get('email')), '')
        if first_email:
            email_el = etree.SubElement(corresp_el, 'email')
            email_el.text = first_email

        # pub-date (use accepted date when available, else today)
        now = datetime.now()
        pub_day = accepted['day'] if accepted else str(now.day).zfill(2)
        pub_month = accepted['month'] if accepted else str(now.month).zfill(2)
        pub_year = accepted['year'] if accepted else str(now.year)

        pd1 = etree.SubElement(am, 'pub-date', {
            'date-type': 'pub', 'publication-format': 'electronic'
        })
        etree.SubElement(pd1, 'day').text = pub_day
        etree.SubElement(pd1, 'month').text = pub_month
        etree.SubElement(pd1, 'year').text = pub_year

        pd2 = etree.SubElement(am, 'pub-date', {
            'date-type': 'collection', 'publication-format': 'electronic'
        })
        etree.SubElement(pd2, 'season').text = 'Jan-Dec'
        etree.SubElement(pd2, 'year').text = pub_year

        if volume:
            etree.SubElement(am, 'volume').text = volume
        if issue:
            etree.SubElement(am, 'issue').text = issue
        if elocation_id:
            etree.SubElement(am, 'elocation-id').text = elocation_id

        # history
        if received or accepted:
            hist = etree.SubElement(am, 'history')
            if received:
                dr = etree.SubElement(hist, 'date', {'date-type': 'received'})
                etree.SubElement(dr, 'day').text = received['day']
                etree.SubElement(dr, 'month').text = received['month']
                etree.SubElement(dr, 'year').text = received['year']
            if accepted:
                da = etree.SubElement(hist, 'date', {'date-type': 'accepted'})
                etree.SubElement(da, 'day').text = accepted['day']
                etree.SubElement(da, 'month').text = accepted['month']
                etree.SubElement(da, 'year').text = accepted['year']

        # abstract (Spanish)
        if abstract_es:
            abs_el = etree.SubElement(am, 'abstract')
            etree.SubElement(abs_el, 'title').text = 'Resumen'
            for para in abstract_es.split('\n\n'):
                if para.strip():
                    etree.SubElement(abs_el, 'p').text = para.strip()

        # trans-abstract (English)
        if abstract_en:
            ta = etree.SubElement(am, 'trans-abstract', {f'{{{self.XML_NS}}}lang': 'en'})
            etree.SubElement(ta, 'title').text = 'Abstract'
            for para in abstract_en.split('\n\n'):
                if para.strip():
                    etree.SubElement(ta, 'p').text = para.strip()

        # kwd-group Spanish
        if keywords_es:
            kg_es = etree.SubElement(am, 'kwd-group', {f'{{{self.XML_NS}}}lang': 'es'})
            etree.SubElement(kg_es, 'title').text = 'Palabras clave:'
            for kw in keywords_es:
                etree.SubElement(kg_es, 'kwd').text = kw

        # kwd-group English
        if keywords_en:
            kg_en = etree.SubElement(am, 'kwd-group', {f'{{{self.XML_NS}}}lang': 'en'})
            etree.SubElement(kg_en, 'title').text = 'Keywords:'
            for kw in keywords_en:
                etree.SubElement(kg_en, 'kwd').text = kw

        # counts
        counts = etree.SubElement(am, 'counts')
        etree.SubElement(counts, 'fig-count', {'count': str(fig_count)})
        etree.SubElement(counts, 'table-count', {'count': str(table_count)})
        etree.SubElement(counts, 'equation-count', {'count': '0'})
        etree.SubElement(counts, 'ref-count', {'count': str(len(references))})

        # ---- Body -------------------------------------------------------
        body = etree.SubElement(article, 'body')

        # Build citation index once so every paragraph can resolve in-text refs
        citation_index = self._build_citation_index(references)
        media_index = self._build_media_index(body_sections)

        intro_labels = {'introducción', 'introduction', 'introduccion'}
        methods_labels = {'metodología', 'methodology', 'metodos', 'methods', 'método'}
        results_labels = {'resultados', 'results'}
        discussion_labels = {'discusión', 'discussion', 'discusion'}
        rd_labels = {'resultados y discusión', 'results and discussion', 'resultados y discusion'}
        concl_labels = {'conclusión', 'conclusion', 'conclusiones', 'conclusions'}

        for section_title, items in body_sections:
            tl = (section_title or '').strip().lower()
            sec_attrs: Dict[str, str] = {}
            if tl in intro_labels:
                sec_attrs['sec-type'] = 'intro'
            elif tl in methods_labels:
                sec_attrs['sec-type'] = 'methods'
            elif tl in results_labels:
                sec_attrs['sec-type'] = 'results'
            elif tl in discussion_labels:
                sec_attrs['sec-type'] = 'discussion'
            elif tl in rd_labels or ('resultados' in tl and 'discusi' in tl):
                sec_attrs['sec-type'] = 'results|discussion'
            elif tl in concl_labels:
                sec_attrs['sec-type'] = 'conclusions'

            sec = etree.SubElement(body, 'sec', sec_attrs)
            etree.SubElement(sec, 'title').text = section_title

            for item in items:
                if isinstance(item, str):
                    self._add_para_with_xrefs(sec, item, citation_index, media_index)
                elif isinstance(item, dict):
                    p_wrap = etree.SubElement(sec, 'p')
                    label = (item.get('label') or '').strip()
                    caption = (item.get('caption') or '').strip()
                    filename = (item.get('filename') or '').strip()
                    item_attrs: Dict[str, str] = {}
                    if item.get('rid'):
                        item_attrs['id'] = item['rid']

                    # Both figures and tables are represented as <fig> with external href.
                    media_el = etree.SubElement(p_wrap, 'fig', item_attrs)

                    if label:
                        etree.SubElement(media_el, 'label').text = label
                    if caption:
                        cap_el = etree.SubElement(media_el, 'caption')
                        etree.SubElement(cap_el, 'title').text = caption
                    if filename:
                        etree.SubElement(media_el, 'graphic', {f'{{{self.XLINK_NS}}}href': filename})

        if len(body) == 0:
            sec = etree.SubElement(body, 'sec')
            etree.SubElement(sec, 'title').text = 'Contenido'
            etree.SubElement(sec, 'p').text = 'Documento convertido.'

        # ---- Back / references ------------------------------------------
        if references:
            back = etree.SubElement(article, 'back')
            rl = etree.SubElement(back, 'ref-list')
            etree.SubElement(rl, 'title').text = 'Referencias Bibliográficas'
            for i, ref_text in enumerate(references, 1):
                ref = etree.SubElement(rl, 'ref', {'id': f'B{i}'})
                self._build_mixed_citation_with_links(ref, ref_text)
                self._build_element_citation(ref, ref_text)

        return article


class JATSToDocxConverter(BaseConverter):
    """Convert JATS XML to Microsoft Word document"""
    
    @property
    def source_formats(self) -> List[str]:
        return ['xml']
    
    @property
    def target_formats(self) -> List[str]:
        return ['docx']
    
    def convert(self, input_path: str, output_path: str) -> bool:
        """
        Convert JATS XML to DOCX
        Creates a formatted Word document from JATS structure
        """
        try:
            self.ensure_directory(output_path)
            
            # Parse XML
            tree = etree.parse(input_path)
            root = tree.getroot()
            
            # Check if it's JATS (look for article element)
            if root.tag not in ['article', '{http://jats.nlm.nih.gov}article']:
                raise ConversionError("El archivo XML no es formato JATS válido")
            
            # Create Word document
            doc = Document()
            
            # Extract and add content
            self._add_front_matter(doc, root)
            self._add_body(doc, root)
            self._add_references(doc, root)
            
            # Save document
            doc.save(output_path)
            
            return True
            
        except etree.XMLSyntaxError as e:
            raise ConversionError(f"Error de sintaxis XML: {str(e)}")
        except Exception as e:
            raise ConversionError(f"Conversión JATS XML a DOCX falló: {str(e)}")
    
    def _add_front_matter(self, doc: Document, root: etree.Element):
        """Add title, authors, and abstract to document"""
        
        # Find front element
        front = root.find('.//front')
        if front is None:
            return
        
        # Title
        title_elem = front.find('.//article-title')
        if title_elem is not None and title_elem.text:
            doc.add_heading(title_elem.text, level=1)
        
        # Authors
        authors = front.findall('.//contrib[@contrib-type="author"]')
        if authors:
            author_names = []
            for author in authors:
                surname = author.find('.//surname')
                given_names = author.find('.//given-names')
                if surname is not None:
                    name = surname.text or ''
                    if given_names is not None and given_names.text:
                        name = f"{given_names.text} {name}"
                    author_names.append(name)
            
            if author_names:
                doc.add_paragraph(', '.join(author_names))
        
        # Abstract
        abstract = front.find('.//abstract')
        if abstract is not None:
            doc.add_heading('Resumen', level=2)
            for p in abstract.findall('.//p'):
                if p.text:
                    doc.add_paragraph(p.text.strip())
        
        # Keywords
        kwd_groups = front.findall('.//kwd-group')
        if kwd_groups:
            doc.add_heading('Palabras clave', level=2)
            keywords = []
            for kwd_group in kwd_groups:
                for kwd in kwd_group.findall('.//kwd'):
                    if kwd.text:
                        keywords.append(kwd.text)
            if keywords:
                doc.add_paragraph(', '.join(keywords))
    
    def _add_body(self, doc: Document, root: etree.Element):
        """Add body sections to document"""
        
        body = root.find('.//body')
        if body is None:
            return
        
        for sec in body.findall('.//sec'):
            # Section title
            title = sec.find('.//title')
            if title is not None and title.text:
                doc.add_heading(title.text, level=2)
            
            # Section paragraphs
            for p in sec.findall('.//p'):
                text = ''.join(p.itertext()).strip()
                if text:
                    doc.add_paragraph(text)
    
    def _add_references(self, doc: Document, root: etree.Element):
        """Add references to document"""
        
        back = root.find('.//back')
        if back is None:
            return
        
        ref_list = back.find('.//ref-list')
        if ref_list is None:
            return
        
        doc.add_heading('Referencias', level=2)
        
        for ref in ref_list.findall('.//ref'):
            mixed_citation = ref.find('.//mixed-citation')
            if mixed_citation is not None:
                text = ''.join(mixed_citation.itertext()).strip()
                if text:
                    doc.add_paragraph(text, style='List Number')
