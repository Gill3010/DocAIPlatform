"""
Procesador de documentos Word para DocAI Platform
"""
import os
from docx import Document
from docx.shared import Inches
import xml.etree.ElementTree as ET
from xml.dom import minidom
import json

class DocumentProcessor:
    def __init__(self):
        self.supported_formats = ['.docx', '.doc']

    def extract_content(self, file_path):
        """Alias de extract_text_from_docx para compatibilidad."""
        return self.extract_text_from_docx(file_path)
    
    def extract_text_from_docx(self, file_path):
        """Extrae texto de un archivo Word"""
        try:
            doc = Document(file_path)
            content = {
                'paragraphs': [],
                'tables': [],
                'metadata': {}
            }
            
            # Extraer párrafos
            for para in doc.paragraphs:
                text = (para.text or "").strip()
                if text:
                    content['paragraphs'].append({
                        'text': text,
                        'style': para.style.name if para.style else 'Normal'
                    })
            
            # Extraer tablas
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                content['tables'].append(table_data)
            
            # Metadata básica
            content['metadata'] = {
                'filename': os.path.basename(file_path),
                'paragraphs_count': len(content['paragraphs']),
                'tables_count': len(content['tables'])
            }
            
            return content
            
        except Exception as e:
            raise Exception(f"Error procesando documento: {str(e)}")
    
    def convert_to_xml(self, content, output_path):
        """Convierte el contenido extraído a XML"""
        try:
            # Crear elemento raíz
            root = ET.Element("document")
            
            # Metadata
            metadata = ET.SubElement(root, "metadata")
            for key, value in content['metadata'].items():
                meta_elem = ET.SubElement(metadata, key)
                meta_elem.text = str(value)
            
            # Párrafos
            paragraphs = ET.SubElement(root, "paragraphs")
            for i, para in enumerate(content['paragraphs']):
                para_elem = ET.SubElement(paragraphs, "paragraph", id=str(i))
                para_elem.set("style", para['style'])
                para_elem.text = para['text']
            
            # Tablas
            if content['tables']:
                tables = ET.SubElement(root, "tables")
                for i, table in enumerate(content['tables']):
                    table_elem = ET.SubElement(tables, "table", id=str(i))
                    for j, row in enumerate(table):
                        row_elem = ET.SubElement(table_elem, "row", id=str(j))
                        for k, cell in enumerate(row):
                            cell_elem = ET.SubElement(row_elem, "cell", id=str(k))
                            cell_elem.text = cell
            
            # Formatear XML
            xml_str = minidom.parseString(ET.tostring(root)).toprettyxml(indent="  ")
            
            # Guardar archivo
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(xml_str)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Error convirtiendo a XML: {str(e)}")
    
    def process_document(self, input_path, output_dir):
        """Procesa un documento completo"""
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"Archivo no encontrado: {input_path}")
        
        # Verificar formato
        file_ext = os.path.splitext(input_path)[1].lower()
        if file_ext not in self.supported_formats:
            raise ValueError(f"Formato no soportado: {file_ext}")
        
        # Extraer contenido
        content = self.extract_text_from_docx(input_path)
        
        # Crear nombre de archivo de salida
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        output_path = os.path.join(output_dir, f"{base_name}.xml")
        
        # Convertir a XML
        xml_path = self.convert_to_xml(content, output_path)
        
        return {
            'input_file': input_path,
            'output_file': xml_path,
            'content_summary': {
                'paragraphs': len(content['paragraphs']),
                'tables': len(content['tables']),
                'total_text_length': sum(len(p['text']) for p in content['paragraphs'])
            }
        }
