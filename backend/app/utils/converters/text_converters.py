"""
Text and Document Conversion Converters
Handles: TXT ↔ DOCX, PDF ↔ TXT
Preserves document structure: inserts [Imagen] placeholders where images appear (accessibility).
Uses docx2python when available for fuller extraction (headers, footers, footnotes).
"""
import re
from docx import Document
from docx.oxml.ns import qn
from typing import List, Tuple

from app.utils.base_converter import BaseConverter, ConversionError


class TextToDocxConverter(BaseConverter):
    """Convert plain text to DOCX"""
    
    @property
    def source_formats(self) -> List[str]:
        return ['txt']
    
    @property
    def target_formats(self) -> List[str]:
        return ['docx']
    
    def convert(self, input_path: str, output_path: str) -> bool:
        """Convert text file to DOCX"""
        try:
            self.ensure_directory(output_path)
            
            # Read text file
            with open(input_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            # Create DOCX
            doc = Document()
            
            # Add paragraphs
            for paragraph in text_content.split('\n\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            
            doc.save(output_path)
            return True
        except Exception as e:
            raise ConversionError(f"Text to DOCX conversion failed: {str(e)}")


def _iter_docx_blocks(doc):
    """Iterate paragraphs and table cells in document order. Yields (block, is_table_cell)."""
    from docx.document import Document as _Document
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph as DocxParagraph

    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield (DocxParagraph(child, doc), False)
        elif child.tag == qn("w:tbl"):
            tbl = DocxTable(child, doc)
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        yield (para, True)


def _para_to_text_with_image_placeholders(para) -> str:
    """Extract text from paragraph, inserting [Imagen] where inline images appear."""
    parts = []
    for run in para.runs:
        has_image = any(
            "blip" in str(desc.tag).lower()
            for desc in run._element.iterdescendants()
        )
        if has_image:
            parts.append("[Imagen]")
        if run.text:
            parts.append(run.text)
    return "".join(parts)


def _docx_to_text_docx2python(input_path: str, output_path: str) -> bool:
    """Use docx2python for fuller extraction (headers, footers, footnotes, document). Normalizes image placeholders to [Imagen]."""
    try:
        from docx2python import docx2python
        with docx2python(input_path) as r:
            text = r.text or ''
            # Normalize image placeholders (docx2python: ----imageN.ext----) to [Imagen]
            text = re.sub(r'----image\d+[^>]*----', '[Imagen]', text, flags=re.IGNORECASE)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text.strip())
            return True
    except ImportError:
        return False
    except Exception:
        return False


class DocxToTextConverter(BaseConverter):
    """Convert DOCX to plain text. Uses docx2python (headers, footers, footnotes) when available, else python-docx. Inserts [Imagen] where images appear."""

    @property
    def source_formats(self) -> List[str]:
        return ['docx']

    @property
    def target_formats(self) -> List[str]:
        return ['txt']

    def convert(self, input_path: str, output_path: str) -> bool:
        """Convert DOCX to text file preserving image positions as [Imagen] placeholders."""
        try:
            self.ensure_directory(output_path)
            if _docx_to_text_docx2python(input_path, output_path):
                return True
            doc = Document(input_path)
            lines = []
            for block, _ in _iter_docx_blocks(doc):
                text = _para_to_text_with_image_placeholders(block).strip()
                if text:
                    lines.append(text)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n\n'.join(lines))
            return True
        except Exception as e:
            raise ConversionError(f"DOCX to text conversion failed: {str(e)}") from e


class PDFToTextConverter(BaseConverter):
    """Extract text from PDF. Inserts [Imagen] where images appear (structure + accessibility)."""

    @property
    def source_formats(self) -> List[str]:
        return ['pdf']

    @property
    def target_formats(self) -> List[str]:
        return ['txt']

    def convert(self, input_path: str, output_path: str) -> bool:
        """Extract text from PDF, inserting [Imagen] placeholders for images in reading order."""
        try:
            self.ensure_directory(output_path)
            import fitz  # PyMuPDF

            all_parts: List[str] = []
            with fitz.open(input_path) as doc:
                for page in doc:
                    elements: List[Tuple[float, float, str]] = []  # (top, left, content)

                    # Text blocks with sort=True for reading order
                    blocks = page.get_text("blocks", sort=True)
                    for b in blocks:
                        x0, y0, x1, y1, text, _bn, block_type = b
                        text_clean = (text or "").strip()
                        if block_type == 1:
                            elements.append((y0, x0, "[Imagen]"))
                        elif text_clean:
                            elements.append((y0, x0, text_clean))

                    # Embedded images (text blocks may miss some)
                    for img in page.get_images(full=True):
                        for r in page.get_image_rects(img[0]):
                            elements.append((r.y0, r.x0, "[Imagen]"))

                    elements.sort(key=lambda e: (e[0], e[1]))
                    # Dedupe consecutive [Imagen] from same area
                    prev_was_img = False
                    for _, _, content in elements:
                        if content == "[Imagen]":
                            if prev_was_img:
                                continue
                            prev_was_img = True
                        else:
                            prev_was_img = False
                        all_parts.append(content)
                    all_parts.append("")

            text_out = "\n\n".join(p.strip() for p in all_parts if p.strip())
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text_out)
            return True
        except ImportError as e:
            raise ConversionError(
                "PDF → texto requiere PyMuPDF. Instala con: pip install PyMuPDF"
            ) from e
        except Exception as e:
            raise ConversionError(f"PDF to text conversion failed: {str(e)}") from e
